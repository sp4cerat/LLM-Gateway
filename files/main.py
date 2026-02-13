"""
LLM Gateway v1.5 - Cost-Optimized AI Routing with Tool Calling
=============================================
Universal LLM Gateway with Three-Tier Routing, Two-Stage Caching,
Budget Control, and OpenAI-Compatible API.

Usage:
    uvicorn main:app --host 0.0.0.0 --port 8000
"""

import os
os.environ["PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK"] = "True"
os.environ["FLAGS_use_mkldnn"] = "0"
import re
import ast
import textwrap
import time
import json
import base64
import asyncio
import httpx
import logging
from collections import deque
from datetime import datetime
from pathlib import Path
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, Request, Depends
from fastapi.responses import JSONResponse, HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware

from models import (
    ChatRequest, ChatResponse, ChatChoice, ChatMessage, UsageInfo,
    RouterAction, CacheInvalidationEvent, GatewayStats
)
from config import load_config, get_api_key, get_gateway_secret, GatewayConfig
from security import (
    check_hard_policy, validate_api_key, generate_request_id,
    ip_limiter, SECURITY_HEADERS, get_cors_origins,
)
from rate_limiter import RateLimiter, BudgetGuard, IdempotencyGuard
from router import IntentRouter
from cache import ExactCache, SemanticCache
from providers import (
    create_provider, LLMProvider, AnthropicProvider,
    STATIC_SYSTEM_PROMPT, DIFF_INSTRUCTION,
    CASCADE_SYSTEM_PROMPT, ESCALATION_MARKER,
    ESCALATION_MARKER_CHEAP_PLUS, ESCALATION_MARKER_MEDIUM, ESCALATION_MARKER_PREMIUM,
    ProviderRateLimitError, ProviderOverloadError,
)
from context import context_budget, output_strategy, estimate_tokens
from metrics import metrics

# v2: Enhanced routing modules
from context_mapper import context_mapper, ContextMap
from enhanced_router import EnhancedRouter, EnhancedRouterResult, RequestContext
from verification import verification_layer
from response_validator import validate_response
from data_collector import get_collector, FreeDataCollector, DataCategory, CollectedData
from tool_executor import (
    TOOL_DEFINITIONS, TOOL_CASCADE_SYSTEM_PROMPT,
    SYNTHESIS_PROMPTS, build_synthesis_prompt,
    extract_blueprint, extract_context_strategy,
    extract_output_architecture, generate_auto_architecture,
    generate_auto_blueprint, execute_tool_calls,
    get_last_source_footer, clear_source_footer,
    tool_web_search,
)
try:
    from web_enrichment import get_web_enricher, classify_needs_enrichment
    _WEB_ENRICHMENT_AVAILABLE = True
except ImportError as e:
    log_boot = logging.getLogger("gateway")
    log_boot.warning(f"web_enrichment not available: {e}")
    _WEB_ENRICHMENT_AVAILABLE = False
    def get_web_enricher(): return None
    def classify_needs_enrichment(q): return {"needs_deep": False, "needs_fact_check": False, "reason": "unavailable"}
from vision_processor import (
    get_vision_pipeline, extract_image_b64, build_text_only_message,
    detect_vision_intent, VisionIntent,
)

# ─── Logging ──────────────────────────────────────────────────────────────────

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(name)s] %(levelname)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
log = logging.getLogger("gateway")

# ─── Global State ─────────────────────────────────────────────────────────────

config: GatewayConfig = None
rate_limiter: RateLimiter = None
budget_guard: BudgetGuard = None
idempotency: IdempotencyGuard = None
intent_router: IntentRouter = None
exact_cache: ExactCache = None
semantic_cache: SemanticCache = None
providers: dict[str, LLMProvider] = {}
gateway_secret: str = ""

# ─── Debug Buffer (last N requests for /debug view) ──────────────────────────
_debug_buffer: deque = deque(maxlen=10)
enhanced_router: EnhancedRouter = None  # v2: multi-layer router
data_collector: FreeDataCollector = None  # Free API data collector (weather, stocks, news)


# ─── Lifespan ─────────────────────────────────────────────────────────────────

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize all gateway components on startup."""
    global config, rate_limiter, budget_guard, idempotency
    global intent_router, exact_cache, semantic_cache, providers, gateway_secret
    global enhanced_router

    log.info("=" * 60)
    log.info("  LLM Gateway v1.5 - Cost-Optimized AI Routing with Tool Calling")
    log.info("=" * 60)

    # Load configuration
    config = load_config()
    log.info(f"Config loaded: strategy={config.routing_strategy}, mock={config.mock_mode}")
    log.info(f"Budget limits: soft=${config.budget.daily_soft_limit}, "
             f"medium=${config.budget.daily_medium_limit}, hard=${config.budget.daily_hard_limit}")

    # Initialize components
    rate_limiter = RateLimiter(config)
    budget_guard = BudgetGuard(config)
    idempotency = IdempotencyGuard()
    intent_router = IntentRouter(config)
    exact_cache = ExactCache()
    semantic_cache = SemanticCache(
        similarity_threshold=config.cache.semantic_similarity_threshold
    )
    gateway_secret = get_gateway_secret()

    # Initialize providers
    _init_providers()

    # v2: Initialize enhanced router (uses intent_router as Layer 4 fallback)
    enhanced_router = EnhancedRouter(config)
    enhanced_router.set_llm_router(intent_router)

    # Initialize free data collector (yfinance, Open-Meteo, NewsAPI)
    global data_collector
    newsapi_key = os.environ.get("NEWSAPI_KEY", "")
    data_collector = get_collector(newsapi_key=newsapi_key)
    log.info(f"Free data collector ready (NewsAPI: {'configured' if newsapi_key else 'not configured'})")

    log.info(f"Gateway ready on {config.host}:{config.port}")
    log.info(f"Mock mode: {'ON' if config.mock_mode else 'OFF'}")
    _cg = config.code_generation
    if _cg.min_tier == "custom" and _cg.custom_model:
        log.info(f"Code generation: custom model → {_cg.custom_model}")
    else:
        log.info(f"Code generation: min_tier={_cg.min_tier}")
    log.info("=" * 60)

    # Start background tasks
    cleanup_task = asyncio.create_task(_periodic_cleanup())

    yield

    # Shutdown
    cleanup_task.cancel()
    log.info("Gateway shutting down")


def _init_providers():
    """Initialize LLM providers from config."""
    global providers

    provider_configs = {
        "router": (config.providers.router_provider, None),
        "cheap": (config.providers.cheap_provider, None),
        "cheap_plus": (config.providers.cheap_plus_provider, None),
        "medium": (config.providers.medium_provider, None),
        "premium": (config.providers.premium_provider, None),
        "embedding": (config.providers.embedding_provider, None),
    }

    for role, (provider_name, _) in provider_configs.items():
        try:
            if config.mock_mode:
                providers[role] = create_provider(provider_name, mock=True)
                log.info(f"  [{role}] Mock provider (no real API calls)")
            else:
                api_key = get_api_key(provider_name)
                providers[role] = create_provider(provider_name, api_key)
                log.info(f"  [{role}] {provider_name} ✓")

                # Set router API key (supports Groq and OpenRouter)
                if role == "router":
                    intent_router.set_api_key(api_key)
        except ValueError as e:
            log.warning(f"  [{role}] {provider_name} - {e} (using mock)")
            providers[role] = create_provider(provider_name, mock=True)


async def _periodic_cleanup():
    """Background task for cache cleanup."""
    while True:
        try:
            await asyncio.sleep(3600)  # Every hour
            cleaned = exact_cache.cleanup_expired()
            if cleaned > 0:
                log.info(f"Cache cleanup: {cleaned} expired entries removed")
        except asyncio.CancelledError:
            break
        except Exception as e:
            log.error(f"Cleanup error: {e}")


# ─── FastAPI App ──────────────────────────────────────────────────────────────

app = FastAPI(
    title="LLM Gateway",
    version="1.5.0",
    description="Cost-Optimized AI Routing Gateway",
    lifespan=lifespan,
)

# ─── Validation Error Handler (log 422 details) ────────────────────────────
from fastapi.exceptions import RequestValidationError

@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    """Log detailed validation errors for debugging 422s."""
    errors = exc.errors()
    log.error(f"422 Validation Error on {request.url.path}: {errors}")
    return JSONResponse(
        status_code=422,
        content={"detail": errors},
    )

app.add_middleware(
    CORSMiddleware,
    allow_origins=get_cors_origins(),
    allow_credentials=True,
    allow_methods=["GET", "POST", "OPTIONS"],
    allow_headers=["Authorization", "X-API-Key", "Content-Type"],
)


# ─── Security Middleware (IP rate limit + headers + body size) ────────────────

@app.middleware("http")
async def security_middleware(request: Request, call_next):
    """IP rate limiting, security headers, and body size check."""
    # Get real IP (behind reverse proxy)
    ip = request.headers.get("X-Forwarded-For", "").split(",")[0].strip()
    if not ip:
        ip = request.headers.get("X-Real-IP", request.client.host if request.client else "unknown")

    # Skip rate limit for health checks
    if request.url.path in ("/health", "/healthz", "/"):
        response = await call_next(request)
        for k, v in SECURITY_HEADERS.items():
            response.headers[k] = v
        return response

    # IP rate limit check
    allowed, reason = ip_limiter.check(ip)
    if not allowed:
        return JSONResponse(
            status_code=429,
            content={"error": reason, "type": "rate_limit"},
            headers={"Retry-After": "60", **SECURITY_HEADERS},
        )

    # Body size check (protect against oversized image uploads)
    content_length = request.headers.get("content-length")
    if content_length and int(content_length) > ip_limiter.max_body_bytes:
        return JSONResponse(
            status_code=413,
            content={"error": f"Request too large. Max {ip_limiter.max_body_bytes // 1_000_000}MB.",
                      "type": "payload_too_large"},
            headers=SECURITY_HEADERS,
        )

    response = await call_next(request)

    # Add security headers to all responses
    for k, v in SECURITY_HEADERS.items():
        response.headers[k] = v

    return response

# Serve static files for dashboard
static_dir = Path(__file__).parent / "static"
if static_dir.exists():
    app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")


# ─── Authentication ───────────────────────────────────────────────────────────

async def verify_auth(request: Request):
    """Verify API key authentication with brute-force protection."""
    if not config or not config.security.require_api_key:
        return True
    if not gateway_secret:
        log.warning("require_api_key=true but GATEWAY_SECRET not set! Allowing request.")
        return True

    # Get client IP for brute-force tracking
    ip = request.headers.get("X-Forwarded-For", "").split(",")[0].strip()
    if not ip:
        ip = request.headers.get("X-Real-IP", request.client.host if request.client else "unknown")

    auth_header = request.headers.get("Authorization", "")
    api_key = request.headers.get("X-API-Key", "")

    provided_key = ""
    if auth_header.startswith("Bearer "):
        provided_key = auth_header[7:]
    elif api_key:
        provided_key = api_key

    if not validate_api_key(provided_key, gateway_secret):
        ip_limiter.record_auth_failure(ip)
        log.warning(f"Auth failed from {ip}")
        raise HTTPException(status_code=401, detail="Invalid API key")

    ip_limiter.record_auth_success(ip)
    return True


# ─── Health & Info Endpoints ──────────────────────────────────────────────────

@app.get("/health")
async def health():
    return {
        "status": "ok",
        "version": "1.5.0",
        "mock_mode": config.mock_mode if config else False,
        "uptime_seconds": round(metrics.get_uptime(), 1),
        "security": {
            "api_key_required": config.security.require_api_key if config else False,
            "policy_gate": config.security.enable_policy_gate if config else False,
        },
    }


@app.get("/security/stats")
async def security_stats(auth: bool = Depends(verify_auth)):
    """IP rate limiter stats (auth required)."""
    stats = ip_limiter.get_stats()
    return {
        "ip_rate_limiter": stats,
        "config": {
            "requests_per_minute": ip_limiter.rpm,
            "requests_per_hour": ip_limiter.rph,
            "burst_limit": ip_limiter.burst_limit,
            "ban_duration_seconds": ip_limiter.ban_duration,
            "max_body_mb": ip_limiter.max_body_bytes // 1_000_000,
        },
    }


@app.get("/v1/models")
async def list_models(auth: bool = Depends(verify_auth)):
    """OpenAI-compatible model listing."""
    models = [
        {"id": "auto", "object": "model", "owned_by": "gateway",
         "description": "Automatic routing (recommended)"},
        {"id": "premium", "object": "model", "owned_by": "gateway",
         "description": f"Premium: {config.providers.premium_model}"},
        {"id": "medium", "object": "model", "owned_by": "gateway",
         "description": f"Medium: {config.providers.medium_model}"},
        {"id": "cheap", "object": "model", "owned_by": "gateway",
         "description": f"Cheap: {config.providers.cheap_model}"},
        {"id": "local", "object": "model", "owned_by": "gateway",
         "description": "Local/fast tier"},
    ]
    return {"object": "list", "data": models}


# ─── Main Chat Endpoint (OpenAI-Compatible) ──────────────────────────────────

async def _llm_call_with_retry(provider, max_retries=1, retry_delay=3.0, **kwargs):
    """Wrap provider.chat() with retry on 429/503."""
    last_error = None
    for attempt in range(1 + max_retries):
        try:
            return await provider.chat(**kwargs)
        except (ProviderRateLimitError, ProviderOverloadError) as e:
            last_error = e
            if attempt < max_retries:
                delay = retry_delay * (attempt + 1)
                log.warning(f"Provider {e.provider}/{e.model}: {type(e).__name__}, "
                            f"retry in {delay:.0f}s (attempt {attempt+1}/{max_retries})")
                await asyncio.sleep(delay)
            else:
                raise


# ─── Code Stitching (continue truncated cheap-model output) ──────────────────

_CODE_STITCH_MAX_ROUNDS = 2   # Max continuation calls (repair step is the fallback)

def _msg_to_dict(msg) -> dict:
    """Convert ChatMessage to OpenAI-compatible dict, preserving tool fields."""
    d = {"role": msg.role, "content": msg.content if msg.content is not None else ""}
    if getattr(msg, 'tool_calls', None):
        d["tool_calls"] = msg.tool_calls
    if getattr(msg, 'tool_call_id', None):
        d["tool_call_id"] = msg.tool_call_id
    if getattr(msg, 'name', None):
        d["name"] = msg.name
    return d


def _has_code_content(content: str) -> bool:
    """
    Check if response contains substantial code.
    Used to decide whether stitching is worth attempting.
    """
    if not content or len(content) < 50:
        return False
    has_code_block = "```" in content
    has_code_indent = any(line.startswith("    ") or line.startswith("\t")
                         for line in content.split("\n")[-20:])
    has_code_syntax = any(kw in content for kw in [
        "def ", "class ", "import ", "function ", "const ", "let ", "var ",
        "return ", "if (", "for (", "while ", "async ", "await ",
        "public ", "private ", "void ", "#include", "package ",
    ])
    return has_code_block or (has_code_indent and has_code_syntax)


def _is_code_truncation(content: str, finish_reason: str) -> bool:
    """
    Detect if a response was truncated mid-code (finish_reason=length).
    For voluntary stops with incomplete code, use _has_code_content instead.
    """
    if finish_reason != "length":
        return False
    return _has_code_content(content)


def _clean_code_response(content: str) -> tuple[str, bool]:
    """
    Clean common formatting artifacts from LLM code responses.

    Gemini Flash sometimes generates diff-style output with +/- prefixes
    on every line, which breaks Python parsing. This strips those markers
    while preserving the actual code.

    Returns: (cleaned_content, was_modified)
    """
    if not content or "```" not in content:
        return content, False

    # Find code blocks and check for diff markers
    modified = False
    result_parts = []
    last_end = 0

    for m in re.finditer(r'(```\w*\n)(.*?)(\n```)', content, re.DOTALL):
        # Add text before this code block
        result_parts.append(content[last_end:m.start()])

        fence_open = m.group(1)
        code = m.group(2)
        fence_close = m.group(3)

        # Check if this looks like a diff (most lines start with + or -)
        lines = code.split("\n")
        diff_lines = sum(1 for l in lines if l.startswith("+") or l.startswith("-"))
        total_nonblank = sum(1 for l in lines if l.strip())

        if total_nonblank > 5 and diff_lines / max(total_nonblank, 1) > 0.5:
            # More than 50% diff-prefixed lines → strip markers
            cleaned_lines = []
            for line in lines:
                if line.startswith("+") and not line.startswith("+++"):
                    cleaned_lines.append(line[1:])  # Strip leading +
                elif line.startswith("-") and not line.startswith("---"):
                    continue  # Remove deleted lines from diff
                else:
                    cleaned_lines.append(line)
            code = "\n".join(cleaned_lines)
            modified = True
            log.info(f"Cleaned diff-style code: {diff_lines}/{total_nonblank} "
                     f"diff lines stripped")

        result_parts.append(fence_open + code + fence_close)
        last_end = m.end()

    if not modified:
        # Also check for unclosed code blocks with diff markers
        unclosed = re.search(r'(```\w*\n)(.*?)$', content[last_end:], re.DOTALL)
        if unclosed:
            code = unclosed.group(2)
            lines = code.split("\n")
            diff_lines = sum(1 for l in lines if l.startswith("+") or l.startswith("-"))
            total_nonblank = sum(1 for l in lines if l.strip())

            if total_nonblank > 5 and diff_lines / max(total_nonblank, 1) > 0.5:
                cleaned_lines = []
                for line in lines:
                    if line.startswith("+") and not line.startswith("+++"):
                        cleaned_lines.append(line[1:])
                    elif line.startswith("-") and not line.startswith("---"):
                        continue
                    else:
                        cleaned_lines.append(line)

                result_parts.append(content[last_end:unclosed.start() + last_end])
                result_parts.append(unclosed.group(1) + "\n".join(cleaned_lines))
                modified = True
                log.info(f"Cleaned diff-style code (unclosed block): "
                         f"{diff_lines}/{total_nonblank} diff lines stripped")
                return "".join(result_parts), True

    if modified:
        result_parts.append(content[last_end:])
        return "".join(result_parts), True

    return content, False


def _extract_code_from_response(content: str) -> tuple[str, str, str]:
    """
    Extract the LARGEST Python code block from a markdown-wrapped response.

    Returns: (prefix, code, suffix)
    - prefix: everything before the code (including the opening fence)
    - code: the actual code content
    - suffix: the closing fence and anything after
    """
    # Find all code blocks (including unclosed ones)
    blocks = []
    for m in re.finditer(r'(```(?:python|py|python3)?\n)(.*?)(\n```|\Z)', content, re.DOTALL):
        blocks.append({
            "start": m.start(),
            "end": m.end(),
            "prefix_fence": m.group(1),
            "code": m.group(2),
            "suffix_fence": m.group(3),
        })

    if not blocks:
        # No code fences — check if the whole response is code
        if any(kw in content for kw in ["def ", "class ", "import "]):
            return "", content, ""
        return content, "", ""

    # Use the largest code block (most likely the main implementation)
    best = max(blocks, key=lambda b: len(b["code"]))
    prefix = content[:best["start"]] + best["prefix_fence"]
    suffix = best["suffix_fence"] + content[best["end"]:]
    return prefix, best["code"], suffix


def _trim_broken_tail(content: str) -> tuple[str, int]:
    """
    Find the last point where code is syntactically valid Python and trim there.
    Also fixes unclosed code fences.

    This fixes the core stitch problem: Flash stops mid-line/mid-block,
    leaving unclosed brackets/strings. Stitching from a broken point fails.
    By trimming back to the last valid parse point, continuation starts clean.

    Returns: (trimmed_content, lines_removed)
    """
    # ── Fix 1: Close unclosed code fences ──
    fence_pattern = re.compile(r'^```', re.MULTILINE)
    fences = fence_pattern.findall(content)
    if len(fences) % 2 != 0:
        # Unclosed fence — close it so the model sees complete markdown
        content = content.rstrip() + "\n```\n"
        log.info("Trim broken tail: closed unclosed code fence")

    prefix, code, suffix = _extract_code_from_response(content)
    if not code:
        return content, 0

    lines = code.split("\n")
    if len(lines) < 5:
        return content, 0

    # ── Fix 2: Trim to last valid Python parse point ──
    try:
        ast.parse(textwrap.dedent(code))
        return content, 0
    except SyntaxError:
        pass

    # Linear scan from the end (usually only 1-15 lines need trimming)
    for trim in range(1, min(30, len(lines) - 5)):
        candidate = "\n".join(lines[:-trim])
        try:
            ast.parse(textwrap.dedent(candidate))
            result = prefix + candidate + "\n```\n"
            log.info(f"Trim broken tail: removed last {trim} lines "
                     f"({len(lines)} → {len(lines) - trim}), code now parses OK")
            return result, trim
        except SyntaxError:
            continue

    # Couldn't find a valid parse point in the last 30 lines — give up
    log.warning(f"Trim broken tail: no valid parse point found in last 30 lines")
    return content, 0


def _stitch_code_chunks(chunks: list[str]) -> str:
    """
    Intelligently stitch code continuation chunks.
    Handles overlap where the model repeats the last few lines.
    """
    if not chunks:
        return ""
    if len(chunks) == 1:
        return chunks[0]

    result = chunks[0]
    for chunk in chunks[1:]:
        if not chunk.strip():
            continue

        # Try to find overlap: the model sometimes repeats the last 1-5 lines
        result_lines = result.rstrip().split("\n")
        chunk_lines = chunk.lstrip().split("\n")

        best_overlap = 0
        # Check last 10 lines of result against first 10 of chunk
        for overlap_len in range(1, min(11, len(result_lines), len(chunk_lines))):
            tail = "\n".join(result_lines[-overlap_len:]).strip()
            head = "\n".join(chunk_lines[:overlap_len]).strip()
            if tail == head:
                best_overlap = overlap_len

        if best_overlap > 0:
            # Remove overlapping lines from chunk
            chunk_trimmed = "\n".join(chunk_lines[best_overlap:])
            result = result.rstrip() + "\n" + chunk_trimmed
            log.info(f"Code stitch: removed {best_overlap} overlapping lines")
        else:
            # No overlap detected — just concatenate
            result = result.rstrip() + "\n" + chunk.lstrip()

    return result


async def _try_code_stitching(
    request_id: str,
    provider,
    model: str,
    messages: list,
    system_prompt: str,
    partial_content: str,
    max_tokens: int,
    temperature: float,
    reason: str = "hard_truncation",
) -> tuple[str, list[float], list[int], bool]:
    """
    Attempt to continue a truncated/incomplete code response by calling the cheap model again.

    Args:
        reason: "hard_truncation" (finish_reason=length) or "incomplete_code" (validator detected)

    Returns:
        (stitched_content, costs, tokens, success)
    """
    costs = []
    tokens = []

    # ── Smart trim: cut back to last valid Python parse point ──
    # Flash often stops mid-line/mid-block leaving unclosed brackets/strings.
    # Stitching from a broken point perpetuates the error.
    # Trim back to the last syntactically valid point first.
    trimmed_content, lines_removed = _trim_broken_tail(partial_content)
    if lines_removed > 0:
        log.info(f"[{request_id}] Smart trim: removed {lines_removed} broken "
                 f"lines before stitching "
                 f"({len(partial_content)} → {len(trimmed_content)} chars)")

    chunks = [trimmed_content]

    # Different continuation prompts depending on reason
    if reason == "hard_truncation":
        continuation_prompt = (
            "Dein vorheriger Output wurde abgeschnitten (max_tokens erreicht). "
            "Setze EXAKT dort fort wo du aufgehört hast. "
            "Wiederhole NICHT was du bereits geschrieben hast. "
            "Schreibe NUR den fehlenden Rest. "
            "Falls der Code bereits vollständig war, schreibe nur: COMPLETE"
        )
    else:  # incomplete_code
        continuation_prompt = (
            "Dein Code oben ist unvollständig — es fehlen noch Teile "
            "(fehlende Klassen, Funktionen, if __name__ Block, etc.). "
            "Schreibe NUR den fehlenden Rest des Codes. "
            "Wiederhole NICHT was du bereits geschrieben hast. "
            "Falls der Code doch bereits vollständig ist, schreibe nur: COMPLETE"
        )

    for round_num in range(1, _CODE_STITCH_MAX_ROUNDS + 1):
        current_content = _stitch_code_chunks(chunks)

        # Build continuation messages:
        # Original messages + assistant's partial response + "continue" instruction
        continuation_msgs = list(messages)  # Copy original
        continuation_msgs.append(
            ChatMessage(role="assistant", content=current_content)
        )
        continuation_msgs.append(
            ChatMessage(role="user", content=continuation_prompt)
        )

        try:
            result = await _llm_call_with_retry(
                provider,
                messages=continuation_msgs,
                model=model,
                max_tokens=max_tokens,
                temperature=temperature,
                system_prompt=system_prompt,
                use_cache=False,
            )
        except Exception as e:
            log.warning(f"[{request_id}] Code stitch round {round_num} failed: {e}")
            break

        costs.append(result.get("cost_usd", 0))
        tokens.append(result["usage"].get("total_tokens", 0))
        continuation = result.get("content", "").strip()
        finish = result.get("finish_reason", "stop")

        # Clean diff markers if present
        continuation, _ = _clean_code_response(continuation)

        log.info(f"[{request_id}] Code stitch round {round_num}: "
                 f"cont_len={len(continuation)}, finish={finish}, "
                 f"cost=${result.get('cost_usd', 0):.6f}")

        # Model says it's done
        if "COMPLETE" in continuation and len(continuation) < 50:
            log.info(f"[{request_id}] Code stitch: model reports COMPLETE "
                     f"after {round_num} rounds")
            return _stitch_code_chunks(chunks), costs, tokens, True

        if not continuation or len(continuation) < 10:
            log.info(f"[{request_id}] Code stitch: empty continuation, stopping")
            break

        chunks.append(continuation)

        # If this continuation finished normally → we're done
        if finish != "length":
            log.info(f"[{request_id}] Code stitch: finished normally "
                     f"after {round_num} rounds, total_len={len(_stitch_code_chunks(chunks))}")
            return _stitch_code_chunks(chunks), costs, tokens, True

    # Exhausted rounds or error — return what we have
    stitched = _stitch_code_chunks(chunks)
    success = len(chunks) > 1  # At least one continuation succeeded
    log.info(f"[{request_id}] Code stitch: {'partial' if success else 'failed'} "
             f"after {len(chunks)-1} rounds, total_len={len(stitched)}")
    return stitched, costs, tokens, success


async def _try_code_repair(
    request_id: str,
    provider,
    model: str,
    messages: list,
    system_prompt: str,
    broken_content: str,
    validation_errors: list[str],
    max_tokens: int,
    temperature: float,
) -> tuple[str, list[float], list[int], bool]:
    """
    Ask the cheap model to FIX its own broken code instead of escalating to premium.

    Unlike stitching (which continues), repair sends the full broken code back
    with the validator's error details and asks for a corrected version.

    Returns:
        (repaired_content, costs, tokens, success)
    """
    # Build error summary for the model
    error_summary = "; ".join(validation_errors[:5])  # Max 5 errors

    repair_msgs = list(messages)  # Original conversation
    repair_msgs.append(
        ChatMessage(role="assistant", content=broken_content)
    )
    repair_msgs.append(
        ChatMessage(
            role="user",
            content=(
                f"Dein Code hat Fehler die automatisch erkannt wurden:\n"
                f"→ {error_summary}\n\n"
                f"Bitte schreibe den KOMPLETTEN korrigierten Code. "
                f"Achte besonders auf: alle Klammern schließen, "
                f"alle Code-Blöcke mit ``` schließen, "
                f"vollständige Klassen/Funktionen, if __name__ Block. "
                f"Schreibe den gesamten Code nochmal korrekt."
            ),
        )
    )

    try:
        result = await _llm_call_with_retry(
            provider,
            messages=repair_msgs,
            model=model,
            max_tokens=max_tokens,
            temperature=temperature,
            system_prompt=system_prompt,
            use_cache=False,
        )
    except Exception as e:
        log.warning(f"[{request_id}] Code repair call failed: {e}")
        return broken_content, [], [], False

    cost = result.get("cost_usd", 0)
    tok = result.get("usage", {}).get("total_tokens", 0)
    repaired = result.get("content", "").strip()

    # Clean diff markers if present
    repaired, was_cleaned = _clean_code_response(repaired)
    if was_cleaned:
        log.info(f"[{request_id}] Code repair: cleaned diff markers from repair output")

    if not repaired or len(repaired) < 50:
        log.warning(f"[{request_id}] Code repair: empty/short response ({len(repaired)} chars)")
        return broken_content, [cost], [tok], False

    log.info(f"[{request_id}] Code repair: got {len(repaired)} chars "
             f"(original {len(broken_content)}), cost=${cost:.6f}")
    return repaired, [cost], [tok], True


def _filter_source_footer(llm_response: str, source_footer: str) -> tuple[str, str]:
    """Filter source footer to only include sources actually referenced in the LLM response.
    Returns (updated_response, filtered_footer) with consistent renumbering.

    Matches these reference formats in the LLM response:
      - [1], [2], [3]                  (new format)
      - [1, 3, 5]                      (grouped new format)
      - (Quelle 1), (Quelle 1, 2, 3)  (legacy format)
      - (Source 1, 3)                  (legacy English)

    The source footer contains lines like "[1] https://... — Title".
    Only keeps footer lines whose [N] number was actually referenced.
    If no numbered references are found, returns all sources unchanged.
    Renumbers both the response text AND footer to be sequential.
    """
    import re

    referenced_nums = set()

    # Pattern 1: [N] inline references — e.g. "text [1]." or "text [3]"
    # Match [N] that is NOT at the start of a line (to avoid matching footer lines)
    for m in re.finditer(r'(?<!^)\[(\d+)\]', llm_response, re.MULTILINE):
        referenced_nums.add(int(m.group(1)))

    # Pattern 2: [N, M, ...] grouped references — e.g. "[1, 3, 5]"
    for m in re.finditer(r'\[(\d+(?:\s*,\s*\d+)+)\]', llm_response):
        for num_str in re.findall(r'\d+', m.group(1)):
            referenced_nums.add(int(num_str))

    # Pattern 3: (Quelle N) or (Quelle N, M) — legacy format
    for m in re.findall(
        r'\((?:Quelle|Source|Src|Q)[\s:]+([0-9,\s]+)\)', llm_response, re.IGNORECASE
    ):
        for num_str in re.findall(r'\d+', m):
            referenced_nums.add(int(num_str))

    # If no numbered references found at all, the model didn't use any sources
    # This usually means the sources were irrelevant — don't return footer
    if not referenced_nums:
        return llm_response, ""

    # Parse footer lines and keep only referenced ones, building renumber map
    footer_lines = source_footer.strip().split("\n")
    filtered = []
    renumber_map = {}  # old_num → new_num
    new_idx = 1
    for line in footer_lines:
        # Keep header lines (---, "Quellen:", "Sources:")
        if line.startswith("---") or line.endswith(":"):
            filtered.append(line)
            continue
        # Parse [N] from footer line
        num_match = re.match(r'\[(\d+)\]', line)
        if num_match:
            orig_num = int(num_match.group(1))
            if orig_num in referenced_nums:
                renumber_map[orig_num] = new_idx
                filtered.append(re.sub(r'^\[\d+\]', f'[{new_idx}]', line))
                new_idx += 1

    # Return filtered footer (or all if filtering removed everything)
    if new_idx > 1 and renumber_map:
        # Renumber references in the response text to match new footer numbering
        updated_response = llm_response
        # Use placeholders to avoid collision (e.g., [4]→[2] then [2]→[1])
        # Step 1: Replace all old refs with unique placeholders
        for old_num in renumber_map:
            updated_response = updated_response.replace(
                f"[{old_num}]", f"[__REF_{old_num}__]")
        # Step 2: Replace placeholders with new numbers
        for old_num, new_num in renumber_map.items():
            updated_response = updated_response.replace(
                f"[__REF_{old_num}__]", f"[{new_num}]")
        return updated_response, "\n".join(filtered)
    return llm_response, source_footer


@app.post("/v1/chat/completions")
async def chat_completions(request: ChatRequest, raw_request: Request = None,
                           auth: bool = Depends(verify_auth)):
    """
    OpenAI-compatible chat completions endpoint.
    Routes requests through the gateway pipeline:
    1. Hard Policy Gate → 2. Rate Limit → 3. Cache → 4. Route → 5. LLM → 6. Cache Store

    Headers:
        X-No-Cache: true  — bypass all caches (for benchmarks/testing)
    """
    request_id = generate_request_id()
    start_time = time.time()

    # Check for cache bypass header (benchmark mode)
    no_cache = False
    if raw_request:
        no_cache = raw_request.headers.get("X-No-Cache", "").lower() in ("true", "1", "yes")

    metrics.increment("requests_total")
    metrics.gauge("active_requests", 1)

    try:
        # Extract user query (last user message) and detect media
        user_query = ""
        has_media = False
        media_types = []
        last_user_idx = -1
        for i, msg in enumerate(reversed(request.messages)):
            if msg.role == "user":
                user_query = msg.text_content
                last_user_idx = len(request.messages) - 1 - i
                if msg.has_media:
                    has_media = True
                    media_types = msg.media_types
                break

        if not user_query and not has_media:
            raise HTTPException(400, "No user message found")

        # Log incoming request details including media
        log.info(f"[{request_id}] Incoming: query='{user_query[:80]}' | "
                 f"has_media={has_media} | media_types={media_types} | "
                 f"msg_count={len(request.messages)}")
        # Debug: log content structure of last user message
        if last_user_idx >= 0:
            last_content = request.messages[last_user_idx].content
            if isinstance(last_content, list):
                block_types = [b.get("type", "?") if isinstance(b, dict) else "text"
                               for b in last_content]
                log.info(f"[{request_id}] Last user msg blocks: {block_types}")
            elif isinstance(last_content, str):
                log.info(f"[{request_id}] Last user msg: str ({len(last_content)} chars)")

        # ── Auto-transcribe audio in messages ─────────────────────────
        # Clients (OpenClaw, etc.) may send audio as content blocks or file refs.
        # Detect, transcribe via Whisper, replace with text before LLM sees it.
        has_audio = has_media and any('audio' in mt.lower() for mt in media_types)

        # Also check for audio file paths in text (OpenClaw pattern)
        # Pattern 1: Files WITH extension (e.g. /path/to/file.ogg)
        _AUDIO_PATH_RE = re.compile(
            r'(/[\w./_-]+\.(?:ogg|opus|mp3|m4a|wav|webm|flac|aac))',
            re.IGNORECASE
        )
        # Pattern 2: OpenClaw [media attached: /path/to/file-without-extension]
        _MEDIA_ATTACHED_RE = re.compile(
            r'\[media attached:\s*(/[\w./_-]+)',
            re.IGNORECASE
        )
        audio_file_path = None
        if not has_audio and user_query:
            # Try extension-based match first
            m = _AUDIO_PATH_RE.search(user_query)
            if m:
                candidate = m.group(1)
                log.info(f"[{request_id}] Audio path candidate (ext): '{candidate}' | "
                         f"exists={os.path.isfile(candidate)}")
                if os.path.isfile(candidate):
                    audio_file_path = candidate
                    has_audio = True

            # Try [media attached:] pattern for extensionless files
            if not has_audio:
                m2 = _MEDIA_ATTACHED_RE.search(user_query)
                if m2:
                    candidate = m2.group(1)
                    # The path might be truncated in the query — try to find the full file
                    if os.path.isfile(candidate):
                        full_path = candidate
                    else:
                        # OpenClaw truncates long UUIDs — glob for partial match
                        import glob
                        matches = glob.glob(candidate + "*")
                        full_path = matches[0] if matches else None

                    if full_path:
                        # Check magic bytes to determine if it's audio
                        is_audio_file = False
                        try:
                            with open(full_path, "rb") as f:
                                header = f.read(12)
                            # ftyp box = MP4/M4A, OggS = OGG, RIFF = WAV, ID3/0xff = MP3, fLaC = FLAC
                            if (b'ftyp' in header[:12] or   # MP4/M4A/3GPP
                                header[:4] == b'OggS' or     # OGG/Opus
                                header[:4] == b'RIFF' or     # WAV
                                header[:3] == b'ID3' or      # MP3 with ID3
                                header[:2] == b'\xff\xfb' or # MP3 without ID3
                                header[:2] == b'\xff\xf3' or # MP3 MPEG2
                                header[:4] == b'fLaC' or     # FLAC
                                header[:4] == b'\x1aE\xdf\xa3'):  # WebM/MKV
                                is_audio_file = True
                        except Exception:
                            pass

                        log.info(f"[{request_id}] Media attached candidate: '{full_path}' | "
                                 f"is_audio={is_audio_file}")
                        if is_audio_file:
                            audio_file_path = full_path
                            has_audio = True

        if has_audio:
            log.info(f"[{request_id}] Audio detected: has_media={has_media} | "
                     f"audio_file_path={audio_file_path} | media_types={media_types}")
            transcribed_text = None
            try:
                # Case 1: Audio content block (input_audio / audio type)
                if isinstance(request.messages[last_user_idx].content, list):
                    for block in request.messages[last_user_idx].content:
                        if isinstance(block, dict):
                            btype = block.get("type", "")
                            if btype == "input_audio":
                                audio_info = block.get("input_audio", {})
                                audio_b64 = audio_info.get("data", "")
                                fmt = audio_info.get("format", "ogg")
                                if audio_b64:
                                    audio_bytes = base64.b64decode(audio_b64)
                                    result = await _transcribe_internal(
                                        audio_bytes, f"audio/{fmt}", "de")
                                    transcribed_text = result.get("text", "")
                                    break
                            elif btype == "audio":
                                audio_b64 = block.get("data", "") or block.get("source", {}).get("data", "")
                                if audio_b64:
                                    audio_bytes = base64.b64decode(audio_b64)
                                    mime = block.get("mime_type", "audio/ogg")
                                    result = await _transcribe_internal(audio_bytes, mime, "de")
                                    transcribed_text = result.get("text", "")
                                    break

                # Case 2: Audio file path on disk (OpenClaw saves to /root/.openclaw/media/)
                if not transcribed_text and audio_file_path:
                    with open(audio_file_path, "rb") as af:
                        audio_bytes = af.read()
                    ext = os.path.splitext(audio_file_path)[1].lower()
                    mime_map = {
                        ".ogg": "audio/ogg", ".opus": "audio/opus", ".mp3": "audio/mpeg",
                        ".m4a": "audio/mp4", ".wav": "audio/wav", ".webm": "audio/webm",
                        ".flac": "audio/flac", ".aac": "audio/aac",
                    }
                    if ext and ext in mime_map:
                        mime = mime_map[ext]
                    else:
                        # No extension — detect MIME from magic bytes
                        header = audio_bytes[:12]
                        if b'ftyp' in header[:12]:
                            mime = "audio/mp4"  # M4A / 3GPP
                        elif header[:4] == b'OggS':
                            mime = "audio/ogg"
                        elif header[:4] == b'RIFF':
                            mime = "audio/wav"
                        elif header[:3] == b'ID3' or header[:2] in (b'\xff\xfb', b'\xff\xf3'):
                            mime = "audio/mpeg"
                        elif header[:4] == b'fLaC':
                            mime = "audio/flac"
                        else:
                            mime = "audio/mp4"  # Safe default for Whisper
                        log.info(f"[{request_id}] Extensionless audio → detected MIME: {mime}")
                    result = await _transcribe_internal(audio_bytes, mime, "de")
                    transcribed_text = result.get("text", "")

                # Case 3: Audio data in media_types but content is base64 somewhere
                if not transcribed_text and has_media:
                    # Try extracting any base64 audio data from content blocks
                    if isinstance(request.messages[last_user_idx].content, list):
                        for block in request.messages[last_user_idx].content:
                            if isinstance(block, dict):
                                # Check for url-style data URIs with audio mime
                                url = (block.get("image_url", {}).get("url", "") or
                                       block.get("url", "") or "")
                                if url.startswith("data:audio/"):
                                    # data:audio/ogg;base64,AAAA...
                                    parts = url.split(",", 1)
                                    if len(parts) == 2:
                                        mime = parts[0].split(":")[1].split(";")[0]
                                        audio_bytes = base64.b64decode(parts[1])
                                        result = await _transcribe_internal(audio_bytes, mime, "de")
                                        transcribed_text = result.get("text", "")
                                        break

                if transcribed_text:
                    # Replace audio message with transcribed text
                    original_text = request.messages[last_user_idx].text_content
                    # Strip [media attached: ...] prefix from original text
                    original_text = re.sub(
                        r'\[media attached:\s*[^\]]*\]\s*', '', original_text
                    ).strip()
                    new_text = f"[Voice message transcription: {transcribed_text}]"
                    if original_text:
                        new_text = f"{original_text}\n\n{new_text}"
                    request.messages[last_user_idx] = ChatMessage(
                        role="user", content=new_text)
                    user_query = new_text
                    has_media = False
                    media_types = []
                    log.info(f"[{request_id}] Audio auto-transcribed: "
                             f"'{transcribed_text[:80]}...' → text-only request")
                else:
                    log.warning(f"[{request_id}] Audio detected but transcription failed/empty")

            except Exception as e:
                log.error(f"[{request_id}] Audio auto-transcription error: {type(e).__name__}: {e}")
                # Fall through — let LLM handle (it'll say it can't process audio)

        # ── Strip images from OLD messages (only keep in latest user msg) ──
        # This prevents sending MB of base64 data from previous turns
        for i, msg in enumerate(request.messages):
            if i != last_user_idx and msg.role == "user" and msg.has_media:
                # Replace multimodal content with text-only
                text = msg.text_content
                img_count = len([b for b in (msg.content if isinstance(msg.content, list) else [])
                                 if isinstance(b, dict) and b.get("type") in ("image_url", "image")])
                if img_count:
                    text = f"{text}\n[{img_count} image(s) from previous message]" if text else f"[{img_count} image(s)]"
                request.messages[i] = ChatMessage(role="user", content=text)
                log.debug(f"Stripped {img_count} image(s) from history message {i}")

        fingerprint = request.fingerprint or ""

        # For multimodal: hash image data into fingerprint so
        # same query + different image = different cache key
        if has_media:
            import hashlib as _hl
            for msg in reversed(request.messages):
                if msg.role == "user" and isinstance(msg.content, list):
                    img_hash_parts = []
                    for block in msg.content:
                        if isinstance(block, dict):
                            url = ""
                            if block.get("type") == "image_url":
                                url = block.get("image_url", {}).get("url", "")
                            elif block.get("type") == "image":
                                url = block.get("source", {}).get("data", "")[:200]
                            if url:
                                # Hash first 2KB of data (fast, sufficient for uniqueness)
                                img_hash_parts.append(
                                    _hl.sha256(url[:2048].encode()).hexdigest()[:16]
                                )
                    if img_hash_parts:
                        fingerprint = f"{fingerprint}|img:{'_'.join(img_hash_parts)}"
                        log.debug(f"[{request_id}] Image fingerprint: {fingerprint[-40:]}")
                    break

        # ─── Detect OpenClaw/Agent sessions ─────────────────────────────
        # OpenClaw sends system prompts containing "default_api", "execute_ipython",
        # "openclaw" etc. These sessions use code-execution patterns that look like
        # hallucinated code to our guards. We must skip those guards for agent sessions.
        # Check ALL messages (system, user, assistant) because OpenClaw may embed
        # its instructions in user messages or conversation history.
        _is_agent_session = False
        _agent_signals = [
            # System prompt signals
            "default_api", "execute_ipython", "openclaw", "openhands",
            "IPythonRunCellAction", "CmdRunAction", "BrowserAction",
            "execute_bash", "<execute_ipython>",
            # Path signals (OpenClaw workspace paths in history)
            "/root/.openclaw/", ".openclaw/workspace",
            # Action format signals
            "OBSERVATION:", "ACTION:", "AgentDelegateAction",
            "FileReadAction", "FileWriteAction",
        ]
        for msg in request.messages:
            if msg.text_content:
                _msg_text = msg.text_content
                if any(sig.lower() in _msg_text.lower() for sig in _agent_signals):
                    _is_agent_session = True
                    _matched = [s for s in _agent_signals if s.lower() in _msg_text.lower()]
                    log.info(f"[{request_id}] Agent session detected "
                             f"(signal '{_matched[0]}' in {msg.role} message)")
                    break
        
        if not _is_agent_session:
            log.info(f"[{request_id}] No agent session detected "
                     f"(checked {len(request.messages)} messages)")
            # Log message summary for debugging agent detection misses
            for i, msg in enumerate(request.messages[:5]):  # First 5 messages
                _preview = msg.text_content[:150].replace('\n', ' ') if msg.text_content else "(empty)"
                log.info(f"[{request_id}] msg[{i}] role={msg.role} len={len(msg.text_content or '')} "
                         f"preview='{_preview}'")
        
        # ─── Step 1: Hard Policy Gate ─────────────────────────────────
        if config.security.enable_policy_gate:
            # Only check the user's actual question, not extracted document content
            # Client puts question first, then \n\n then document text
            policy_text = user_query
            if len(user_query) > 1000:
                # Large content = extracted document appended to question
                # Check only the question part (before first double-newline)
                first_break = user_query.find("\n\n")
                if first_break > 0 and first_break < 500:
                    policy_text = user_query[:first_break]
                else:
                    policy_text = user_query[:500]
                log.debug(f"[{request_id}] Policy check on question only: "
                          f"{len(policy_text)} chars (full: {len(user_query)})")

            violation = check_hard_policy(
                policy_text,
                enable_sensitive=config.security.block_sensitive_data
            )
            if violation:
                metrics.increment("policy_blocks", tags={"category": violation.category})
                raise HTTPException(403, f"Policy violation: {violation.category} - {violation.message}")

        # ─── Step 2: Rate Limiting ────────────────────────────────────
        estimated_tokens_count = estimate_tokens(user_query)
        # Estimate total request size (all messages)
        total_request_tokens = sum(estimate_tokens(m.content) for m in request.messages)

        # ─── Step 3: Idempotency Check ───────────────────────────────
        idem_key = idempotency.get_key(request.messages, request.model, request.temperature or 0.7)
        if not no_cache:
            cached_idem = idempotency.check(idem_key)
            if cached_idem:
                log.info(f"[{request_id}] Idempotency hit")
                return _build_response(request_id, cached_idem, "cached", {"source": "idempotency"},
                                       stream=request.stream)

        # ─── Step 4: Exact Cache ──────────────────────────────────────
        # Image hash is part of fingerprint → same query + different image = cache miss
        # Skip cache for agent sessions — context changes with each tool result
        cache_query = user_query
        if config.cache.exact_cache_enabled and not no_cache and not _is_agent_session:
            exact_hit = exact_cache.get(cache_query, fingerprint)
            if exact_hit:
                latency = (time.time() - start_time) * 1000
                metrics.histogram("request_latency_ms", latency, tags={"tier": "cached"})
                log.info(f"[{request_id}] Exact cache hit ({latency:.0f}ms)"
                         f"{' [multimodal]' if has_media else ''}")
                return _build_response(request_id, exact_hit["response"],
                                       exact_hit.get("model", "cached"),
                                       {"source": "exact_cache", "latency_ms": latency,
                                        "multimodal_cached": has_media},
                                       stream=request.stream)

        # ─── Step 5: Route Intent (v2: Enhanced Multi-Layer Routing) ──
        route_result = None
        cascade_mode = (config.routing_strategy == "cascade")
        size_escalated_from = None
        enhanced_result = None
        context_map = None
        full_context_text = ""
        is_doc_qa = False
        _expects_long_code = False  # Set by code upgrade, used by post-validation
        _needs_web_search = False   # Set by planning/web keyword detection
        tier = None  # Will be set by routing logic below

        if request.model and request.model != "auto":
            # Explicit model selection — skip routing
            tier = request.model if request.model in ("local", "cheap", "cheap_plus", "medium", "premium") else "premium"
            cascade_mode = False

        else:
            # ─── Step 5a: Build Context Map for large requests ────
            # Detect short question + long document (PDF analysis pattern)
            # CRITICAL: Only check the LAST user message length, not total user history.
            # Chat clients (OpenClaw) send many user/assistant message pairs as history.
            # Summing all user messages falsely triggers doc_qa for normal chat.
            query_tokens = estimate_tokens(user_query.split('\n')[0])  # Just the question part
            last_user_msg_tokens = estimate_tokens(user_query)  # Full last user message

            # Doc QA = short question embedded in a SINGLE long user message
            # (e.g. user pastes a 10-page document + asks "summarize this")
            # NOT triggered by: short question + lots of chat history messages
            is_doc_qa = (query_tokens < 50 and last_user_msg_tokens > 800)

            if is_doc_qa:
                log.info(f"[{request_id}] Document QA detected: "
                         f"query={query_tokens} tok, total={total_request_tokens} tok "
                         f"→ skipping enhanced routing")

            if total_request_tokens > 1500 and not is_doc_qa:
                context_parts = [m.text_content for m in request.messages
                                 if m.role in ("system", "user") and len(m.text_content) > 200]
                full_context_text = "\n\n".join(context_parts)
                if len(full_context_text) > 3000:
                    context_map = context_mapper.build(full_context_text)
                    log.info(f"[{request_id}] Context map: {context_map.chunk_count} chunks, "
                             f"~{context_map.total_tokens_est} tok, type={context_map.document_type}")

            # ─── Step 5b: Enhanced routing (large context) ────────
            if context_map and context_map.total_tokens_est > 1500:
                try:
                    request_context = RequestContext(
                        query_tokens=estimated_tokens_count,
                        context_tokens=context_map.total_tokens_est,
                        document_type=context_map.document_type,
                        customer_tier=request.customer_tier or "standard",
                        task_priority=request.task_priority or "normal",
                        budget_utilization=budget_guard.get_status().today_spend / max(0.01, config.budget.daily_hard_limit),
                    )
                    enhanced_result = await enhanced_router.route(
                        query=user_query,
                        context_map=context_map,
                        request_context=request_context,
                    )
                    tier = enhanced_result.action.value
                    cascade_mode = False  # Enhanced router replaces cascade
                    log.info(f"[{request_id}] Enhanced route: {tier} | "
                             f"strategy={enhanced_result.strategy} | "
                             f"layer={enhanced_result.routing_layer}")
                except Exception as e:
                    log.error(f"[{request_id}] Enhanced routing failed: {type(e).__name__}: {e}, "
                              f"falling back to cascade")
                    enhanced_result = None
                    context_map = None  # Prevent chunk retrieval from also failing
                    # Fall through to cascade below (don't use elif)

            # ─── Step 5c: Original cascade (small context) ────────
            if cascade_mode and tier is None:
                # For document QA, use much higher thresholds
                # Gemini Flash has 1M token context — no need to escalate for documents
                if is_doc_qa:
                    CHEAP_MAX_TOKENS = 50000
                    MEDIUM_MAX_TOKENS = 100000
                    routing_tokens = total_request_tokens  # Doc size matters
                    _routing_query = user_query  # No stripping needed for doc QA
                else:
                    CHEAP_MAX_TOKENS = 500
                    MEDIUM_MAX_TOKENS = 3000
                    # Use QUERY complexity for routing, not total request size
                    # OpenClaw/chat clients add 2000-5000 tok of system prompt + history
                    # A simple "Wie wird das Wetter" should stay cheap regardless
                    _routing_query = user_query
                    # Agent sessions: OpenClaw prepends system logs (exec failures, etc.)
                    # to user messages, inflating token count. Extract actual user text.
                    # Pattern: "System: [...] ...\n\n[Telegram Sven ...] actual message"
                    if _is_agent_session and "\n" in _routing_query:
                        # Find the actual Telegram message after system logs
                        _tg_match = re.search(
                            r'\[Telegram [^\]]+\]\s*(.+?)(?:\s*\[message_id:\s*\d+\])?\s*$',
                            _routing_query, re.DOTALL
                        )
                        if _tg_match:
                            _routing_query = _tg_match.group(1).strip()
                            log.debug(f"[{request_id}] Agent routing: stripped system logs, "
                                      f"query='{_routing_query[:80]}'")
                    routing_tokens = estimate_tokens(_routing_query)
                log.info(f"[{request_id}] Cascade routing: routing_tok={routing_tokens} | "
                         f"total_req_tok={total_request_tokens} | "
                         f"thresholds: cheap<{CHEAP_MAX_TOKENS} medium<{MEDIUM_MAX_TOKENS}"
                         f"{' [doc_qa]' if is_doc_qa else ''}")
                _needs_web_search = False
                if routing_tokens >= MEDIUM_MAX_TOKENS:
                    tier = "premium"
                    size_escalated_from = "cheap"
                    log.info(f"[{request_id}] Size-based routing: {routing_tokens} tok → premium")
                elif routing_tokens >= CHEAP_MAX_TOKENS:
                    tier = "medium"
                    size_escalated_from = "cheap"
                    log.info(f"[{request_id}] Size-based routing: {routing_tokens} tok → medium")
                else:
                    tier = "cheap"

                # ── Web search upgrade: cheap → cheap_plus ──
                # Detect queries that need web search even when token count is low.
                if tier == "cheap":
                    _q_lower = _routing_query.lower()
                    _web_search_signals = [
                        # Explicit research requests (DE/EN)
                        "recherchiere", "recherche zu", "recherche über",
                        "research", "investigate", "look up",
                        # Source/data requests
                        "aktuelle studien", "aktuelle quellen", "nutze quellen",
                        "nutze aktuelle", "aktuelle daten", "aktuelle zahlen",
                        "finde heraus", "finde informationen", "suche nach",
                        "search for", "find out about",
                        # Real-time data
                        "aktuelle nachrichten", "latest news",
                        "spielstand", "score of", "ergebnis von",
                        "öffnungszeiten", "opening hours",
                        "wie spät", "what time",
                    ]
                    # Also check: realtime keywords + question words
                    _realtime_kw = ["aktuell", "currently", "gerade", "heute",
                                    "today", "live", "jetzt", "derzeit"]
                    _question_words = ["was", "wie", "what", "how", "wieviel",
                                       "welche", "which", "wann", "when"]
                    _planning_kw = ["plant", "plans", "vorhaben", "position",
                                    "wahlprogramm", "koalition", "regierung",
                                    "gesetzentwurf", "reform", "beschlossen",
                                    "policy", "legislation", "regulation"]
                    _has_realtime = any(kw in _q_lower for kw in _realtime_kw)
                    _has_question = any(w in _q_lower for w in _question_words)
                    _has_planning = any(kw in _q_lower for kw in _planning_kw)

                    needs_web = (
                        any(sig in _q_lower for sig in _web_search_signals)
                        or (_has_realtime and _has_question)
                        or (_has_planning and _has_question)
                    )
                    # Exclude queries handled by dedicated tools (weather, stocks)
                    _tool_handled = ["dax", "aktie", "aktien", "börse", "stock",
                                     "share price", "kurs", "portfolio", "index",
                                     "dow", "nasdaq", "s&p", "nikkei", "ftse",
                                     "wetter", "weather", "temperatur", "temperature",
                                     "forecast", "vorhersage"]
                    if needs_web and any(t in _q_lower for t in _tool_handled):
                        needs_web = False
                        log.info(f"[{request_id}] Web upgrade suppressed: "
                                 f"query matches tool-handled pattern, staying on cheap")
                    if needs_web:
                        tier = "cheap_plus"
                        _needs_web_search = True
                        log.info(f"[{request_id}] Web search upgrade: cheap → cheap_plus "
                                 f"(query needs web data)")

                # ── Code generation upgrade: cheap → medium ──
                # Only for LONG code tasks (200+ lines, full programs).
                # Benchmark v3 shows Gemini 2.0 Flash is better at short code (96% vs 94%),
                # but Gemini 3 Flash is more reliable for long code (no indent/diff bugs).
                # Gemini 3 Flash ($0.50/$3.00/M) is better for complex code tasks.
                if tier == "cheap":
                    _q_lower = _routing_query.lower()
                    _long_code_signals = [
                        # Explicit long/full program requests (DE)
                        "vollständiges programm", "vollständiges skript",
                        "vollständigen code", "vollständige implementierung",
                        "vollständiges python", "komplettes programm",
                        # Explicit long/full program requests (EN)
                        "full implementation", "complete program",
                        "complete implementation", "full application",
                        "entire program", "entire application",
                    ]
                    _line_count_signals = [
                        "mindestens 100 zeilen", "mindestens 150 zeilen",
                        "mindestens 200 zeilen", "mindestens 250 zeilen",
                        "at least 100 lines", "at least 150 lines",
                        "at least 200 lines",
                    ]
                    is_long_code = (
                        any(sig in _q_lower for sig in _long_code_signals)
                        or any(sig in _q_lower for sig in _line_count_signals)
                    )
                    if is_long_code:
                        tier = "medium"
                        _expects_long_code = True
                        log.info(f"[{request_id}] Code upgrade: cheap → medium "
                                 f"(long code generation detected, same price)")

            # ─── Step 5d: Classic intent routing (fallback) ───────
            if tier is None:
                route_result = await intent_router.route(user_query)
                tier = route_result.action.value
                if tier == "cache_only":
                    return _build_response(
                        request_id,
                        {"content": "Could you please be more specific? Your request is too vague for a meaningful answer."},
                        "gateway",
                        {"source": "cache_only", "reason": route_result.reason},
                        stream=request.stream
                    )

        # ─── Step 5e: Code generation escalation ─────────────────────────
        # Escalate to better model when code generation is detected.
        # Detection layers (language-independent):
        # 1. EnhancedRouter/IntentRouter LLM sets is_code_generation=True
        # 2. LLM code detection call for small-context queries (test console)
        # 3. Agent sessions: recent tool_calls for code tools (exec, write, sub_agent)
        # NOTE: response_type="code_suggestion" is NOT used as signal — it fires
        #       for questions ABOUT code too (false positives like "webroot" queries).
        _agent_custom_model = None
        if tier in ("cheap", "local"):
            _code_cfg = config.code_generation
            
            # Check 1: LLM-classified code generation
            _is_code = False
            if enhanced_result and enhanced_result.is_code_generation:
                _is_code = True
            elif route_result and getattr(route_result, 'is_code_generation', False):
                _is_code = True
            
            # Check 2: LLM code detection (when no router ran yet)
            # This happens for small-context cascade queries (test console, simple API).
            # Groq/OpenRouter call: ~50ms, ~0.001ct — negligible cost.
            # ONLY trust the explicit is_code_generation flag, NOT response_type.
            # (response_type=code_suggestion also fires for questions ABOUT code,
            #  e.g. "Ist in deinem workspace ein webroot" → false positive)
            if not _is_code and not enhanced_result and not route_result:
                try:
                    _code_check = await intent_router.route(user_query)
                    route_result = _code_check  # Save for metadata display
                    if _code_check.is_code_generation:
                        _is_code = True
                        log.info(f"[{request_id}] LLM code detection: is_code=True "
                                 f"(action={_code_check.action.value}, "
                                 f"rt={_code_check.response_type})")
                    else:
                        log.debug(f"[{request_id}] LLM code detection: is_code=False "
                                  f"(action={_code_check.action.value}, "
                                  f"rt={_code_check.response_type})")
                except Exception as e:
                    log.debug(f"[{request_id}] LLM code detection failed: {e}")
            
            # Check 3: Agent sessions — recent tool_calls for code tools
            # Only check last 4 messages (current exchange), not deep history,
            # to avoid false positives when user switched topics after coding.
            # SKIP when last message is a tool result: the tier was already decided
            # on the previous call; this is just processing output (e.g. exec ls).
            _last_is_tool_result = (
                request.messages and request.messages[-1].role == "tool"
            )
            if not _is_code and _is_agent_session and not _last_is_tool_result:
                _code_tools = set(_code_cfg.code_tool_names)
                for msg in reversed(request.messages[-4:]):
                    if getattr(msg, 'tool_calls', None):
                        for tc in msg.tool_calls:
                            _tc_name = tc.get("function", {}).get("name", "")
                            if _tc_name in _code_tools:
                                _is_code = True
                                break
                    if _is_code:
                        break
            
            if _is_code:
                _old_tier = tier
                _source = (
                    "is_code_generation" if (
                        (enhanced_result and enhanced_result.is_code_generation)
                        or (route_result and getattr(route_result, 'is_code_generation', False))
                    ) else "agent_tool_history"
                )
                if _code_cfg.min_tier == "custom" and _code_cfg.custom_model:
                    _agent_custom_model = _code_cfg.custom_model
                    tier = "medium"  # Use medium provider slot but override model
                    log.info(f"[{request_id}] Code escalation: {_old_tier} → "
                             f"custom model ({_agent_custom_model}) [{_source}]")
                elif _code_cfg.min_tier in ("medium", "premium"):
                    tier = _code_cfg.min_tier
                    log.info(f"[{request_id}] Code escalation: {_old_tier} → "
                             f"{tier} (config min_tier={_code_cfg.min_tier}) [{_source}]")
                # "cheap" min_tier = no escalation

        # ─── Step 5f: Vision Pipeline (images → preprocess → route) ─────
        # Safety net: if no routing path set tier, default to cheap
        if tier is None:
            log.warning(f"[{request_id}] No routing path set tier, defaulting to cheap")
            tier = "cheap"

        vision_result = None
        if has_media:
            # Extract image data from last user message
            last_user_msg = None
            for msg in reversed(request.messages):
                if msg.role == "user" and msg.has_media:
                    last_user_msg = msg
                    break

            image_b64 = extract_image_b64(last_user_msg.content) if last_user_msg else None

            if image_b64:
                pipeline = get_vision_pipeline()
                vision_result = await pipeline.process(image_b64, user_query)
                strategy = vision_result["strategy"]
                intent = vision_result["intent"]

                log.info(f"[{request_id}] Vision pipeline: strategy={strategy} | "
                         f"intent={intent.value} | "
                         f"can_skip={vision_result['can_skip_image']}")

                if vision_result["can_skip_image"] and vision_result["text_description"]:
                    # ── OCR/local vision sufficient → text-only, stay cheap ──
                    text_content = build_text_only_message(
                        last_user_msg.content,
                        vision_result["text_description"]
                    )
                    # Replace multimodal message with text-only
                    for i, msg in enumerate(request.messages):
                        if msg is last_user_msg:
                            request.messages[i] = ChatMessage(
                                role="user", content=text_content
                            )
                            break
                    has_media = False  # No longer multimodal
                    # Recalculate tokens (image is gone now!)
                    total_request_tokens = sum(
                        estimate_tokens(m.content) for m in request.messages
                    )
                    # Re-evaluate tier: image tokens inflated the count earlier,
                    # now with text-only it may fit in cheap tier again
                    if size_escalated_from and cascade_mode:
                        CHEAP_MAX_TOKENS = 500
                        if total_request_tokens < CHEAP_MAX_TOKENS:
                            log.info(f"[{request_id}] Vision: post-OCR tokens "
                                     f"({total_request_tokens}) < {CHEAP_MAX_TOKENS} "
                                     f"→ downgrading {tier}→cheap")
                            tier = "cheap"
                            size_escalated_from = None
                        elif total_request_tokens < 3000 and tier == "premium":
                            tier = "medium"
                            log.info(f"[{request_id}] Vision: post-OCR → medium")

                    log.info(f"[{request_id}] Vision: image replaced with text "
                             f"({len(text_content)} chars, {total_request_tokens} tok) "
                             f"→ {tier}")
                    metrics.increment("vision_image_skipped",
                                      tags={"strategy": strategy})

                elif strategy == "augmented" and vision_result["text_description"]:
                    # ── Partial info → augment + send image to medium ──
                    # Prepend description to the text part of the message
                    if isinstance(last_user_msg.content, list):
                        augment_block = {
                            "type": "text",
                            "text": vision_result["text_description"]
                        }
                        # Insert at beginning of content list
                        for i, msg in enumerate(request.messages):
                            if msg is last_user_msg:
                                new_content = [augment_block] + list(msg.content)
                                request.messages[i] = ChatMessage(
                                    role="user", content=new_content
                                )
                                break
                    if tier in ("cheap", "cheap_plus"):
                        previous_tier = tier
                        tier = "medium"
                        size_escalated_from = previous_tier
                        cascade_mode = False
                    log.info(f"[{request_id}] Vision: augmented + image → {tier}")
                    metrics.increment("vision_augmented",
                                      tags={"strategy": strategy})

                else:
                    # ── Passthrough → escalate to medium+ ──
                    if tier in ("cheap", "cheap_plus"):
                        previous_tier = tier
                        tier = "medium"
                        size_escalated_from = previous_tier
                        cascade_mode = False
                    log.info(f"[{request_id}] Vision: passthrough → {tier}")
                    metrics.increment("vision_passthrough")
            else:
                # No extractable image data → simple escalation
                if tier in ("cheap", "cheap_plus"):
                    previous_tier = tier
                    tier = "medium"
                    size_escalated_from = previous_tier
                    cascade_mode = False
                    log.info(f"[{request_id}] Multimodal routing: "
                             f"{previous_tier}→{tier} (no image data)")

        metrics.increment("requests_by_tier", tags={"tier": tier})

        # ─── Step 6: Rate Limit Check (with actual tier) ─────────────
        rate_ok, rate_reason, rate_delay = rate_limiter.check(tier, estimated_tokens_count)
        if not rate_ok:
            metrics.increment("rate_limit_hits")
            raise HTTPException(429, f"Rate limit exceeded: {rate_reason}")

        # ─── Step 7: Budget Check ─────────────────────────────────────
        estimated_cost = _estimate_request_cost(tier, estimated_tokens_count)
        budget_result = budget_guard.check(tier, estimated_cost)
        if not budget_result["allowed"]:
            raise HTTPException(429, f"Budget limit: {budget_result['reason']}")
        if budget_result["delay"] > 0:
            await asyncio.sleep(budget_result["delay"])

        # ─── Step 8: Semantic Cache (if enabled) ──────────────────────
        if config.cache.semantic_cache_enabled and not no_cache and not _is_agent_session and tier in ("cheap", "cheap_plus", "medium", "premium"):
            embedding_fn = None
            if "embedding" in providers:
                embedding_fn = providers["embedding"].get_embedding

            sem_hit = await semantic_cache.get(user_query, fingerprint, embedding_fn)
            if sem_hit:
                latency = (time.time() - start_time) * 1000
                metrics.histogram("request_latency_ms", latency, tags={"tier": "cached"})
                log.info(f"[{request_id}] Semantic cache hit (sim={sem_hit['similarity']:.3f}, {latency:.0f}ms)")
                return _build_response(request_id, sem_hit["response"],
                                       sem_hit.get("model", "cached"),
                                       {"source": "semantic_cache",
                                        "similarity": sem_hit["similarity"],
                                        "latency_ms": latency},
                                       stream=request.stream)

        # ─── Step 8.5: Targeted Chunk Retrieval (v2) ─────────────────
        if enhanced_result and enhanced_result.needed_chunks and context_map and full_context_text:
            targeted_text = context_mapper.retrieve_chunks(
                full_context_text, context_map, enhanced_result.needed_chunks
            )
            if targeted_text:
                # Replace large context with only the relevant chunks
                new_messages = []
                for msg in request.messages:
                    if msg.role == "user" and msg == request.messages[-1]:
                        new_messages.append(ChatMessage(
                            role="user",
                            content=f"RELEVANT CONTEXT:\n{targeted_text}\n\nQUESTION:\n{msg.text_content}"
                        ))
                    else:
                        new_messages.append(msg)
                request_messages_for_llm = new_messages
                log.info(f"[{request_id}] Chunk retrieval: {len(enhanced_result.needed_chunks)} chunks, "
                         f"~{estimate_tokens(targeted_text)} tok (was {context_map.total_tokens_est} tok)")
            else:
                request_messages_for_llm = request.messages
        else:
            request_messages_for_llm = request.messages

        # ─── Step 9: Context Budgeting ────────────────────────────────
        # Use higher budget for doc QA on cheap (Gemini Flash has 1M context)
        budget_tier = tier
        if is_doc_qa and tier == "cheap":
            budget_tier = "cheap_long"
        trimmed_messages, max_output = context_budget.apply(
            budget_tier, request_messages_for_llm, STATIC_SYSTEM_PROMPT if tier == "premium" else ""
        )

        # ─── Step 10: Output Strategy ─────────────────────────────────
        response_type = route_result.response_type if route_result else "explanation_generic"
        out_strategy = output_strategy.get_strategy(
            response_type, tier, request.file_context
        )
        max_output = min(max_output, out_strategy["max_output_tokens"])

        # ─── Step 11: LLM Call (with Tool Calling + Cascade) ──────────
        provider = _get_provider_for_tier(tier)
        model = _get_model_for_tier(tier)
        # Override with custom model for agent code generation
        if _agent_custom_model:
            model = _agent_custom_model
            log.info(f"[{request_id}] Using custom code model: {model}")

        # Build system prompt
        # Skip tool cascade for long content — tool calling is for short queries
        # (e.g. "what's the weather"), not for document analysis. Models like Gemini
        # refuse when they see tool definitions + large text blobs.
        # Check QUERY length, not total request (system prompt + history inflate total)
        # Use last user message tokens — if the user's actual message is long, it's doc analysis.
        # If it's short but total is high due to OpenClaw boilerplate, it's NOT doc analysis.
        user_query_tokens = estimate_tokens(user_query)
        is_document_analysis = (user_query_tokens > 300 or
                                (user_query_tokens > 100 and is_doc_qa))
        if cascade_mode and tier in ("cheap", "cheap_plus") and not size_escalated_from and not is_document_analysis:
            # Tool-aware cascade: cheap model decides what tools to call
            # SKIP for agent sessions — OpenClaw has its own system prompt and tools
            if _is_agent_session:
                # Don't use system_prompt (gets buried under OpenClaw's 16k developer message).
                # Instead inject hints directly into trimmed_messages near the end,
                # right before the last user message, where Gemini will actually see them.
                system_prompt = ""
                from datetime import datetime as _dt_hint
                _today_hint = _dt_hint.now().strftime("%d.%m.%Y")
                _tool_hint = ChatMessage(
                    role="system",
                    content=(
                        f"Heutiges Datum: {_today_hint}\n\n"
                        "CRITICAL TOOL RULES (MUST FOLLOW):\n"
                        "1. NEVER output Python code as text. No print(), no default_api calls as text.\n"
                        "2. For ANY factual question, knowledge, news, politics, events → call web_search tool\n"
                        "3. For weather → call get_weather tool\n"
                        "4. For stock prices → call get_stock_price tool\n"
                        "5. memory_search is ONLY for recalling past user conversations. "
                        "NEVER use memory_search for general knowledge questions.\n"
                        "6. ALWAYS respond in the SAME language as the user's message.\n"
                        "7. WEB SEARCH QUALITY: When calling web_search, use these parameters:\n"
                        '   - search_depth: "medium" (downloads full pages, not just snippets)\n'
                        '   - additional_queries: add 1-2 alternative search terms for better coverage\n'
                        '   - time_filter: Choose based on the query:\n'
                        '     "d" = last 24h → for "heute", "today", breaking news, live events\n'
                        '     "w" = last week → for "diese Woche", current events, recent prices\n'
                        '     "m" = last month → for recent studies, policy changes, new products\n'
                        '     "none" = no limit → for laws, regulations, research, established facts\n'
                        '   Example: web_search(query="Veranstaltungen Konstanz Februar 2026",\n'
                        '            search_depth="medium", time_filter="w",\n'
                        '            additional_queries=["Events Bodensee Wochenende"])\n'
                        '   Example: web_search(query="EU Batterieverordnung 2026",\n'
                        '            search_depth="medium", time_filter="none")'
                        + ("\n⚠️ THIS QUERY REQUIRES CURRENT DATA. You MUST call web_search "
                           "before answering. Do NOT answer from memory alone."
                           if _needs_web_search else "")
                    ),
                )
                # Insert before last user message
                _last_user_idx = None
                for _i in range(len(trimmed_messages) - 1, -1, -1):
                    if trimmed_messages[_i].role == "user":
                        _last_user_idx = _i
                        break
                if _last_user_idx is not None:
                    trimmed_messages.insert(_last_user_idx, _tool_hint)
                else:
                    trimmed_messages.append(_tool_hint)
                use_tools = False
                log.info(f"[{request_id}] Agent session → tool hints injected before last user msg")
            else:
                from datetime import datetime as _dt
                _today_str = _dt.now().strftime("%d.%m.%Y")
                system_prompt = f"Heutiges Datum: {_today_str}\n\n" + TOOL_CASCADE_SYSTEM_PROMPT
        elif tier == "premium":
            system_prompt = STATIC_SYSTEM_PROMPT
        else:
            system_prompt = ""
        use_cache = (tier == "premium")

        # Inject diff instruction if needed
        if out_strategy.get("inject_diff_instruction"):
            if trimmed_messages:
                last_msg = trimmed_messages[-1]
                if last_msg.role == "user":
                    trimmed_messages[-1] = ChatMessage(
                        role="user",
                        content=DIFF_INSTRUCTION + "\n\n" + last_msg.text_content
                    )

        # Pass tool definitions only for cheap cascade mode (short queries)
        # SKIP for agent sessions — OpenClaw has its own tools, gateway tools would interfere
        use_tools = (cascade_mode and tier in ("cheap", "cheap_plus") and not size_escalated_from
                     and not is_document_analysis and not _is_agent_session)
        if is_document_analysis and tier == "cheap":
            log.info(f"[{request_id}] Document analysis detected "
                     f"(query_tok={user_query_tokens}, total={total_request_tokens}) "
                     f"→ skipping tool cascade")
        log.info(f"[{request_id}] Tool decision: use_tools={use_tools} | "
                 f"cascade_mode={cascade_mode} | tier={tier} | "
                 f"size_escalated={size_escalated_from} | "
                 f"query_tok={user_query_tokens} | total_req_tok={total_request_tokens} | "
                 f"is_doc_qa={is_doc_qa} | model={request.model}")

        # Determine which tools to pass to the LLM
        # For agent sessions: merge client tools + gateway tools so the model can
        # use web_search/weather/stocks AND the agent's own tools (exec, write, etc.)
        _gateway_tool_names = set()
        if _is_agent_session and request.tools:
            _gateway_tool_names = {t["function"]["name"] for t in TOOL_DEFINITIONS
                                   if "function" in t}
            # Merge: client tools + gateway tools (skip duplicates)
            _client_tool_names = set()
            for t in request.tools:
                fname = t.get("function", {}).get("name", "")
                if fname:
                    _client_tool_names.add(fname)
            # Add gateway tools that the client doesn't already provide
            _extra_gateway = [t for t in TOOL_DEFINITIONS
                              if t.get("function", {}).get("name") not in _client_tool_names]
            _effective_tools = list(request.tools) + _extra_gateway
            log.info(f"[{request_id}] Agent session: {len(request.tools)} client tools + "
                     f"{len(_extra_gateway)} gateway tools = {len(_effective_tools)} total")
        elif use_tools:
            _effective_tools = TOOL_DEFINITIONS  # Gateway tools only
        else:
            _effective_tools = None

        # Log message structure for agent session debugging
        if _is_agent_session:
            _msg_summary = []
            for i, m in enumerate(trimmed_messages):
                _has_tc = bool(getattr(m, 'tool_calls', None))
                _has_tcid = bool(getattr(m, 'tool_call_id', None))
                _clen = len(m.text_content) if m.text_content else 0
                _msg_summary.append(f"{m.role}({'tc' if _has_tc else ''}"
                                    f"{'tid' if _has_tcid else ''}"
                                    f",{_clen}c)")
            log.info(f"[{request_id}] Agent messages: {' → '.join(_msg_summary)}")

        llm_result = await _llm_call_with_retry(
            provider,
            messages=trimmed_messages,
            model=model,
            max_tokens=max_output,
            temperature=request.temperature or 0.7,
            system_prompt=system_prompt,
            use_cache=use_cache,
            tools=_effective_tools,
            tool_choice=request.tool_choice if _is_agent_session else None,
        )

        # Track cascade costs
        cascade_costs = [llm_result.get("cost_usd", 0)]
        cascade_tokens = [llm_result["usage"].get("total_tokens", 0)]
        escalated_to = None
        validation_escalated = None  # Set by Step 11d if response validation fails
        code_stitched = False        # Set by Step 11d if code stitching succeeded
        code_repaired = False        # Set by Step 11d if code repair succeeded
        tools_used = []
        _agent_tool_calls = None     # Set for agent sessions to pass through
        web_enrichment_ctx = ""  # Set by web enrichment if pre-search was done
        _search_depth = "snippets"   # Set by synthesis step
        _analysis_mode = "factual"   # Set by synthesis step
        _ctx_mode = "full"           # Set by context strategy
        
        log.info(f"[{request_id}] LLM result: tool_calls={llm_result.get('tool_calls') is not None} | "
                 f"content_len={len(llm_result.get('content',''))} | "
                 f"finish={llm_result.get('finish_reason','?')} | "
                 f"input_tok={llm_result['usage'].get('prompt_tokens',0)}")

        # ─── Step 11b: Tool Calling Loop + Escalation ─────────────────
        if use_tools or _is_agent_session:
            tool_calls = llm_result.get("tool_calls")
            response_content = llm_result.get("content", "").strip()

            if tool_calls and _is_agent_session:
                # ── Agent session: split tool_calls into gateway vs client tools ──
                _gw_calls = [tc for tc in tool_calls 
                             if tc.get("function", {}).get("name") in _gateway_tool_names]
                _client_calls = [tc for tc in tool_calls
                                 if tc.get("function", {}).get("name") not in _gateway_tool_names]
                
                if _client_calls:
                    _agent_tool_calls = _client_calls
                    log.info(f"[{request_id}] Agent client tool_calls (pass-through): "
                             f"{[tc.get('function',{}).get('name','?') for tc in _client_calls]}")
                    # Gemini sometimes duplicates tool calls as plaintext in content
                    # e.g. content="print(default_api.memory_search(...))" + tool_calls=[memory_search]
                    # Strip the plaintext code so OpenClaw doesn't show it to the user
                    if response_content and any(p in response_content for p in 
                            ["default_api.", "print(", ".memory_search(", ".web_search(",
                             ".exec(", ".write(", ".read("]):
                        log.info(f"[{request_id}] Stripping plaintext code from agent content: "
                                 f"'{response_content[:100]}'")
                        response_content = ""
                        llm_result["content"] = ""  # Keep in sync
                
                if _gw_calls:
                    # Gateway tools need server-side execution
                    tool_calls = _gw_calls
                    log.info(f"[{request_id}] Agent gateway tools (executing): "
                             f"{[tc.get('function',{}).get('name','?') for tc in _gw_calls]}")
                else:
                    # Only client tools — skip gateway execution
                    tool_calls = None
                    _final_content = response_content
                    _source_footer = ""
                    _filtered_footer = ""
            
            # ── Agent forced web_search: detect hallucinated/lazy responses ──
            # Three triggers:
            # 1. _needs_web_search=True but model didn't call web_search
            # 2. Model response contains citation markers ([1], [2], "Quellen:")
            #    but no web_search was called → hallucinated sources
            # 3. Model deflects ("konnte keine Informationen finden", "empfehle 
            #    lokale Webseite") without actually searching → lazy deflection
            _forced_web_search = False
            if _is_agent_session and not tool_calls:
                _has_web_call = False
                if llm_result.get("tool_calls"):
                    _has_web_call = any(
                        tc.get("function", {}).get("name") == "web_search"
                        for tc in llm_result["tool_calls"]
                    )
                _resp = response_content or ""
                _resp_lower = _resp.lower()
                # Detect hallucinated source patterns in response
                import re as _re_hall
                _has_fake_citations = bool(
                    _re_hall.search(r'\[\d+\]', _resp)  # [1], [2], etc.
                    and ("http" in _resp or "quellen" in _resp_lower
                         or "quelle" in _resp_lower or "source" in _resp_lower)
                )

                # Detect lazy deflection: model says "I couldn't find anything"
                # without ever calling web_search
                _deflection_phrases = [
                    "konnte keine", "keine spezifischen",
                    "keine informationen", "keine ergebnisse",
                    "keine veranstaltungen", "keine events",
                    "keine daten", "nichts gefunden",
                    "couldn't find", "no results", "no information",
                    "could not find", "unable to find",
                ]
                _deflection_signals = [
                    "empfehle", "empfehlung", "webseite", "website",
                    "kalender", "nachrichten", "prüfen",
                    "recommend", "check the", "visit the",
                ]
                _has_deflection = (
                    any(p in _resp_lower for p in _deflection_phrases)
                    and any(s in _resp_lower for s in _deflection_signals)
                    and not _has_web_call
                )

                # Check if query is a short greeting/chat (not worth web searching)
                _rq_words = _routing_query.split()
                _question_words_check = {"was", "wie", "what", "how", "wieviel",
                                         "welche", "which", "wann", "when", "wo",
                                         "where", "wer", "who", "warum", "why"}
                _is_greeting = (
                    len(_rq_words) <= 2
                    and not any(w.lower() in _question_words_check
                                for w in _rq_words)
                )

                _should_force = (
                    not _is_greeting and not _has_web_call
                    and (_needs_web_search or _has_fake_citations or _has_deflection)
                )

                if _is_greeting and _has_fake_citations:
                    # Short greeting with hallucinated sources → strip citations
                    # instead of forcing a web search for "Huhu" or "Hallo"
                    _resp = _re_hall.sub(r'\[\d+\]', '', _resp)
                    # Strip Quellen/Sources footer
                    _resp = _re_hall.split(
                        r'\n+(?:Quellen|Sources|Quelle|Source)\s*:',
                        _resp
                    )[0].strip()
                    response_content = _resp
                    llm_result["content"] = _resp
                    log.info(f"[{request_id}] Stripped hallucinated citations "
                             f"from greeting response (query='{_routing_query}')")

                elif _should_force:
                    _trigger = ("needs_web" if _needs_web_search 
                                else "hallucinated_citations" if _has_fake_citations
                                else "lazy_deflection")
                    log.warning(f"[{request_id}] Agent forced web_search "
                                f"(trigger={_trigger}): model didn't call web_search, "
                                f"executing for query: '{_routing_query[:80]}'")
                    metrics.increment("agent_forced_web_search",
                                      tags={"trigger": _trigger})
                    # Build a synthetic web_search tool call
                    import uuid as _uuid
                    _forced_tc_id = f"forced_{_uuid.uuid4().hex[:8]}"
                    tool_calls = [{
                        "id": _forced_tc_id,
                        "type": "function",
                        "function": {
                            "name": "web_search",
                            "arguments": json.dumps({
                                "query": _routing_query,
                                "search_depth": "medium",
                            }),
                        },
                    }]
                    _forced_web_search = True
                    response_content = _resp

            # Execute gateway tool calls (normal flow)
            if tool_calls:
                # ── Model requested tool calls → execute free APIs ──
                log.info(f"[{request_id}] Tool calls: "
                         f"{[tc.get('function',{}).get('name','?') for tc in tool_calls]}")

                tool_results = await execute_tool_calls(tool_calls)
                tools_used = [tr["name"] for tr in tool_results]
                metrics.increment("tool_calls_executed",
                                  tags={"tools": ",".join(tools_used)})

                # Build follow-up messages with tool results (OpenAI format)
                follow_up = []

                # ── Extract analysis parameters from Round 1 ──
                _search_depth = "deep"  # Minimum: tool_executor always upgrades snippets→deep
                _analysis_mode = "factual"
                for tc in tool_calls:
                    if tc.get("function", {}).get("name") == "web_search":
                        try:
                            _tc_args = json.loads(tc.get("function", {}).get("arguments", "{}"))
                            _raw_depth = _tc_args.get("search_depth",
                                            _tc_args.get("depth", "deep"))
                            # Normalize: min/medium/max → snippets/deep/thorough
                            _depth_map = {"min": "deep", "medium": "deep", "max": "thorough",
                                          "snippets": "deep"}
                            _search_depth = _depth_map.get(_raw_depth, "deep")
                            _analysis_mode = _tc_args.get("analysis_mode", "factual")
                        except (json.JSONDecodeError, AttributeError):
                            pass

                # Extract output_architecture JSON from Round 1 assistant content
                _output_arch = extract_output_architecture(response_content) if response_content else {}
                _blueprint = extract_blueprint(response_content) if (response_content and not _output_arch) else ""

                # Auto-architecture fallback: if Round 1 generated no architecture/blueprint
                # but we have deep/thorough depth, generate a reasonable default
                if not _output_arch and not _blueprint and _search_depth in ("deep", "thorough"):
                    # Get original user query for context
                    _user_query = ""
                    for _m in reversed(trimmed_messages):
                        if _m.role == "user":
                            _user_query = _m.content[:200]
                            break
                    _output_arch = generate_auto_architecture(_user_query, _analysis_mode)
                    log.info(f"[{request_id}] Auto-architecture generated "
                             f"(R1 content_len={len(response_content or '')})")

                # ── Extract context strategy (JSON or legacy text) ──
                _ctx_strategy = extract_context_strategy(
                    response_content) if response_content else {"mode": "full"}

                # Build dynamic synthesis prompt
                synthesis_prompt = build_synthesis_prompt(
                    depth=_search_depth,
                    analysis_mode=_analysis_mode,
                    blueprint=_blueprint,
                    output_architecture=_output_arch or None,
                )

                # ── Agent sessions: use lighter synthesis prompt ──
                # The full synthesis prompt (400+ lines) is designed for the test console.
                # For agent/Telegram context, use a concise prompt with structure guidance.
                if _is_agent_session and "web_search" in tools_used:
                    from datetime import datetime as _dt
                    _today = _dt.now().strftime("%d.%m.%Y")
                    synthesis_prompt = (
                        f"Du bist ein hilfreicher Recherche-Assistent. Heutiges Datum: {_today}.\n"
                        "Dir wurden Suchergebnisse bereitgestellt.\n\n"
                        "ANTWORT-FORMAT:\n"
                        "- Beginne mit einer klaren Zusammenfassung in 1-2 Sätzen\n"
                        "- Gliedere mit **fetten Überschriften** wo sinnvoll (2-4 Abschnitte)\n"
                        "- Verwende konkrete Zahlen und Fakten aus den Quellen\n"
                        "- Nenne Datumsangaben und zeitliche Einordnung wenn relevant\n"
                        "- Schließe mit einem kurzen **Fazit** oder Einordnung ab\n"
                        "- Liste Quellen als nummerierte Links am Ende: [1] URL — Titel\n\n"
                        "REGELN:\n"
                        "- SPRACHE: Antworte in der Sprache der Nutzerfrage (Deutsch → Deutsch)\n"
                        "- Extrahiere Fakten aus den Quellen, nicht aus deinem Wissen\n"
                        "- Bei Widersprüchen: benenne beide Positionen\n"
                        "- Keine Code-Blöcke, keine API-Aufrufe\n"
                        "- Halte die Antwort kompakt (Telegram-Kontext), aber informativ\n\n"
                        "SELBST-BEWERTUNG (PFLICHT):\n"
                        "Füge am ENDE deiner Antwort (nach den Quellen) EXAKT diesen Block ein:\n"
                        '<!--QUALITY:{"answered":true/false,"confidence":"high"/"medium"/"low",'
                        '"source_count":N,"has_facts":true/false,'
                        '"retry":null oder {"queries":["..."],"time_filter":"d"/"w"/"m"/"none","reason":"..."}}\n'
                        "-->\n"
                        "Feld-Erklärungen:\n"
                        "- answered: true wenn die Frage konkret beantwortet wurde\n"
                        "- confidence: high/medium/low — Qualität der Quellenabdeckung\n"
                        "- source_count: Anzahl genutzter Quellen\n"
                        "- has_facts: true wenn konkrete Fakten/Zahlen/Termine extrahiert\n"
                        "- retry: null wenn Antwort ausreichend. Sonst ein Objekt mit:\n"
                        '  - queries: 1-3 bessere Suchbegriffe (z.B. ["Veranstaltungen Konstanz Februar 2026", "Events Bodensee"])\n'
                        '  - time_filter: "d" (24h), "w" (Woche), "m" (Monat), "none" (unbegrenzt)\n'
                        "  - reason: kurze Begründung warum Nachsuche sinnvoll\n"
                        "Empfehle retry wenn: Quellen irrelevant, zu wenig konkrete Daten, "
                        "falsche Sprache der Ergebnisse, oder Frage nur teilweise beantwortet."
                    )
                    # For agent sessions, use distill context (only question + search results)
                    # instead of sending the full OpenClaw conversation history back
                    _ctx_strategy = {"mode": "distill", "distill_text": ""}
                    log.info(f"[{request_id}] Agent synthesis: using lightweight prompt "
                             f"(depth={_search_depth})")

                follow_up.append({"role": "system", "content": synthesis_prompt})

                # ── Apply Context Strategy: full / recent:N / distill ──
                _ctx_mode = _ctx_strategy.get("mode", "full")
                _original_msgs = len(trimmed_messages)
                _original_tok = sum(len(m.content or "") // 4 for m in trimmed_messages)

                if _ctx_mode == "distill":
                    # DISTILL: Replace full history with compact extract
                    last_user_msg = ""
                    for msg in reversed(trimmed_messages):
                        if msg.role == "user":
                            last_user_msg = msg.text_content or msg.content or ""
                            break
                    _distill_text = _ctx_strategy.get("distill_text", "")
                    if _distill_text:
                        compact_context = (
                            f"Relevanter Kontext aus der Konversation:\n"
                            f"{_distill_text}\n\n"
                            f"Aktuelle Frage:\n{last_user_msg}"
                        )
                    else:
                        # Agent mode: no prior context needed, just the question
                        # Strip OpenClaw system logs if present
                        compact_context = last_user_msg
                        if _is_agent_session and "\n" in compact_context:
                            import re as _re2
                            _tg = _re2.search(
                                r'\[Telegram [^\]]+\]\s*(.+?)(?:\s*\[message_id:\s*\d+\])?\s*$',
                                compact_context, _re2.DOTALL
                            )
                            if _tg:
                                compact_context = _tg.group(1).strip()
                    follow_up.append({"role": "user", "content": compact_context})
                    _r2_msg_count = 1
                    _r2_tok = len(compact_context) // 4
                    log.info(f"[{request_id}] Context DISTILL: "
                             f"{_original_tok} tok → {_r2_tok} tok "
                             f"(saved ~{_original_tok - _r2_tok} tok)")

                elif _ctx_mode == "recent":
                    # RECENT:N: Keep only last N user+assistant pairs
                    n_pairs = _ctx_strategy.get("recent_n", 3)
                    msgs_to_keep = n_pairs * 2  # N pairs = 2N messages
                    system_msgs = [m for m in trimmed_messages if m.role == "system"]
                    non_system = [m for m in trimmed_messages if m.role != "system"]

                    if len(non_system) > msgs_to_keep:
                        trimmed_count = len(non_system) - msgs_to_keep
                        non_system = non_system[-msgs_to_keep:]
                        log.info(f"[{request_id}] Context RECENT:{n_pairs}: "
                                 f"kept {len(non_system)}/{_original_msgs} msgs "
                                 f"(trimmed {trimmed_count})")
                    r2_messages = system_msgs + non_system
                    for msg in r2_messages:
                        follow_up.append(_msg_to_dict(msg))
                    _r2_msg_count = len(r2_messages)

                else:
                    # FULL: Pass entire conversation history (default, safe for code)
                    for msg in trimmed_messages:
                        follow_up.append(_msg_to_dict(msg))
                    _r2_msg_count = _original_msgs

                log.info(f"[{request_id}] Synthesis: depth={_search_depth} "
                         f"analysis={_analysis_mode} "
                         f"arch={'json' if _output_arch else ('legacy' if _blueprint else 'none')} "
                         f"ctx={_ctx_mode} r2_msgs={_r2_msg_count} "
                         f"prompt_len={len(synthesis_prompt)}")

                # Assistant's tool call response (strip internal metadata)
                _r2_content = response_content or ""
                # Remove JSON architecture block — already extracted, waste tokens in R2
                import re as _re
                _r2_content = _re.sub(r'```json\s*\n?\{.*?\}\s*\n?```', '',
                                      _r2_content, flags=_re.DOTALL).strip()
                # Also strip legacy metadata blocks
                for _strip_marker in ["ANALYSE-BLUEPRINT:", "ANALYSIS-BLUEPRINT:",
                                      "KONTEXT-EXTRAKT:", "CONTEXT-EXTRACT:",
                                      "KONTEXT-MODUS:", "CONTEXT-MODE:"]:
                    _strip_idx = _r2_content.find(_strip_marker)
                    if _strip_idx >= 0:
                        _r2_content = _r2_content[:_strip_idx].rstrip()
                assistant_msg = {
                    "role": "assistant",
                    "content": _r2_content if _r2_content else None,
                }
                assistant_msg["tool_calls"] = tool_calls
                follow_up.append(assistant_msg)
                # Tool results
                for tr in tool_results:
                    follow_up.append({
                        "role": "tool",
                        "tool_call_id": tr["tool_call_id"],
                        "content": tr["result"],
                    })
                    log.info(f"[{request_id}] Tool result [{tr['name']}]: "
                             f"{tr['result'][:200]}...")

                # Second LLM call: model formats the tool results into final answer
                # Small delay to avoid OpenRouter rate limits (especially free tier)
                await asyncio.sleep(1.0)
                llm_result = await provider.chat(
                    messages=[],  # Unused when raw_messages is set
                    model=model,
                    max_tokens=max_output,
                    temperature=request.temperature or 0.7,
                    raw_messages=follow_up,
                )
                cascade_costs.append(llm_result.get("cost_usd", 0))
                cascade_tokens.append(llm_result["usage"].get("total_tokens", 0))

                log.info(f"[{request_id}] Tool call complete: {tools_used} | "
                         f"total_cost=${sum(cascade_costs):.6f}")

                # ── Post-synthesis guard: catch JSON/blueprint in synthesis output ──
                _synth_out = llm_result.get("content", "").strip()
                if _synth_out.startswith("{") and _synth_out.endswith("}"):
                    try:
                        _synth_json = json.loads(_synth_out)
                        _synth_keys = set(_synth_json.keys())
                        _blueprint_keys = {"output_architecture", "query", "search_depth",
                                           "context_mode", "context_extract"}
                        if _synth_keys & _blueprint_keys:
                            log.warning(f"[{request_id}] Synthesis returned JSON blueprint "
                                        f"instead of answer (keys: {_synth_keys})")
                            # Retry synthesis with explicit anti-JSON instruction
                            _retry_follow = list(follow_up)
                            _retry_follow.append({
                                "role": "user",
                                "content": (
                                    "FEHLER: Du hast JSON statt einer Antwort ausgegeben. "
                                    "Gib NIEMALS JSON, Code oder Metadaten als Antwort. "
                                    "Beantworte die ursprüngliche Frage als normalen Text "
                                    "mit konkreten Fakten aus den Suchergebnissen oben."
                                ),
                            })
                            _retry_synth = await provider.chat(
                                messages=[], model=model,
                                max_tokens=max_output,
                                temperature=request.temperature or 0.7,
                                raw_messages=_retry_follow,
                            )
                            cascade_costs.append(_retry_synth.get("cost_usd", 0))
                            cascade_tokens.append(
                                _retry_synth["usage"].get("total_tokens", 0))
                            _retry_text = _retry_synth.get("content", "").strip()
                            if _retry_text and not _retry_text.startswith("{"):
                                llm_result = _retry_synth
                                log.info(f"[{request_id}] Synthesis retry succeeded: "
                                         f"{len(_retry_text)} chars")
                            else:
                                log.warning(f"[{request_id}] Synthesis retry also JSON")
                    except (json.JSONDecodeError, TypeError):
                        pass  # Not JSON, that's fine

            # Check for escalation markers (model chose not to use tools)
            elif ESCALATION_MARKER_PREMIUM in response_content:
                escalated_to = "premium"
            elif ESCALATION_MARKER_MEDIUM in response_content:
                escalated_to = "medium"
            elif ESCALATION_MARKER_CHEAP_PLUS in response_content:
                # Legacy marker → try tool-based approach as fallback
                log.info(f"[{request_id}] Legacy CHEAP_PLUS marker → "
                         f"falling back to data_collector")
                collected_data = await data_collector.collect(user_query)
                if collected_data.has_data:
                    enriched_query = data_collector.build_enriched_prompt(
                        user_query, collected_data
                    )
                    enriched_messages = [ChatMessage(role="user", content=enriched_query)]
                    llm_result = await provider.chat(
                        messages=enriched_messages,
                        model=model,
                        max_tokens=max_output,
                        temperature=request.temperature or 0.7,
                        system_prompt="",
                        use_cache=False,
                    )
                    cascade_costs.append(llm_result.get("cost_usd", 0))
                    cascade_tokens.append(llm_result["usage"].get("total_tokens", 0))
                else:
                    escalated_to = "cheap_plus"
            elif ESCALATION_MARKER in response_content:
                escalated_to = "medium"

            # ── Guard: Hallucinated code/JSON instead of tool call ──
            # Sometimes the model outputs Python code (e.g. print(default_api.web_search(...)))
            # or raw JSON blueprints (e.g. {"query": "...", "search_depth": "medium"})
            # or output_architecture JSON instead of actually calling tools.
            # Detection: no tool_calls, response looks like code/JSON, not a real answer.

            # Agent session guard: Gemini sometimes outputs tool calls as plaintext
            # e.g. "print(default_api.memory_search(...))" instead of proper tool_call
            # Strip these patterns and retry
            if _is_agent_session and not tool_calls and response_content:
                _agent_code_patterns = [
                    "default_api.", "print(default_api.", "print(default_api",
                ]
                _rc_stripped = response_content.strip()
                if (any(p in _rc_stripped for p in _agent_code_patterns)
                        and len(_rc_stripped) < 500):
                    log.warning(f"[{request_id}] Agent plaintext tool call detected: "
                                f"'{_rc_stripped[:150]}' → retrying")
                    # Retry with explicit instruction
                    _retry_msgs = list(trimmed_messages)
                    _retry_msgs.append(ChatMessage(
                        role="user",
                        content=(
                            "FEHLER: Du hast Code als Text ausgegeben statt ein Tool zu verwenden. "
                            "Verwende das web_search Tool für Wissensfragen oder antworte direkt. "
                            "Gib NIEMALS Python-Code wie print(default_api...) aus."
                        ),
                    ))
                    llm_result = await _llm_call_with_retry(
                        provider, messages=_retry_msgs, model=model,
                        max_tokens=max_output,
                        temperature=request.temperature or 0.7,
                        system_prompt=system_prompt, use_cache=False,
                        tools=_effective_tools,
                        tool_choice=request.tool_choice if _is_agent_session else None,
                    )
                    cascade_costs.append(llm_result.get("cost_usd", 0))
                    cascade_tokens.append(llm_result["usage"].get("total_tokens", 0))
                    tool_calls = llm_result.get("tool_calls")
                    response_content = llm_result.get("content", "").strip()
                    # Re-check for agent tool routing
                    if tool_calls and _is_agent_session:
                        _gw_calls = [tc for tc in tool_calls
                                     if tc.get("function", {}).get("name") in _gateway_tool_names]
                        _client_calls = [tc for tc in tool_calls
                                         if tc.get("function", {}).get("name") not in _gateway_tool_names]
                        if _client_calls:
                            _agent_tool_calls = _client_calls
                        if _gw_calls:
                            tool_calls = _gw_calls
                        else:
                            tool_calls = None
                            _final_content = response_content
                            _source_footer = ""
                            _filtered_footer = ""

            # SKIP remaining hallucination guard for agent sessions
            if not tool_calls and not escalated_to and response_content and not _is_agent_session:
                _hallucination_patterns = [
                    "default_api.", "print(", ".web_search(", ".memory_search(",
                    ".get_weather(", ".get_stock_price(", ".get_news(",
                    "import requests", "requests.get(", "requests.post(",
                    "api.search(", "api.query(",
                ]
                _rc_lower = response_content.lower().strip()
                _rc_stripped = response_content.strip()
                _is_hallucinated_code = (
                    any(p.lower() in _rc_lower for p in _hallucination_patterns)
                    and len(response_content) < 500  # Real answers are longer
                    and ("(" in response_content and ")" in response_content)
                )

                # Also detect raw JSON output (model dumps blueprint instead of tool call)
                # Patterns: {"query": "...", "search_depth": "..."} 
                #           {"output_architecture": {...}}
                _is_hallucinated_json = False
                if _rc_stripped.startswith("{") and _rc_stripped.endswith("}"):
                    try:
                        _parsed_json = json.loads(_rc_stripped)
                        _json_keys = set(_parsed_json.keys())
                        # Known internal-only JSON patterns that should never be user-facing
                        _internal_json_keys = {
                            "query", "search_depth", "output_architecture",
                            "context_mode", "context_extract", "tool_call",
                        }
                        if _json_keys & _internal_json_keys:
                            _is_hallucinated_json = True
                            log.warning(
                                f"[{request_id}] Hallucinated JSON detected "
                                f"(keys: {_json_keys}): '{_rc_stripped[:120]}'")
                    except (json.JSONDecodeError, TypeError):
                        pass

                if _is_hallucinated_code or _is_hallucinated_json:
                    _hall_type = "JSON blueprint" if _is_hallucinated_json else "code"
                    log.warning(
                        f"[{request_id}] Hallucinated {_hall_type} detected: "
                        f"'{response_content[:100]}' → retrying with reinforced prompt")
                    metrics.increment("hallucinated_code_retries")

                    # ── Fast path: if hallucinated JSON contains a query, execute it directly ──
                    _direct_synthesis_done = False
                    if _is_hallucinated_json:
                        try:
                            _hall_json = json.loads(_rc_stripped)
                            _hall_query = _hall_json.get("query", "")
                            _hall_depth = _hall_json.get("search_depth", "medium")
                            _hall_arch = _hall_json.get("output_architecture")
                            
                            # If no query in JSON, use the user's original question
                            if not _hall_query and _hall_arch:
                                _hall_query = user_query
                                log.info(f"[{request_id}] No query in JSON, using user query: "
                                         f"'{_hall_query[:80]}'")
                            
                            if _hall_query:
                                log.info(f"[{request_id}] Extracting query from hallucinated JSON: "
                                         f"'{_hall_query}' (depth={_hall_depth})")
                                # Map depth
                                _depth_map = {"min": "snippets", "medium": "deep", "max": "thorough"}
                                _actual_depth = _depth_map.get(_hall_depth, "deep")
                                # Execute web search directly
                                _ws_result = await tool_web_search(
                                    query=_hall_query, depth=_actual_depth)
                                if _ws_result and not _ws_result.startswith("Keine Ergebnisse"):
                                    # Build synthesis directly
                                    _synth_prompt = build_synthesis_prompt(
                                        _actual_depth,
                                        output_architecture=_hall_arch or {},
                                    )
                                    _synth_msgs = []
                                    for m in trimmed_messages:
                                        _synth_msgs.append(
                                            {"role": m.role, "content": m.text_content})
                                    _synth_msgs.append({
                                        "role": "user",
                                        "content": f"{_ws_result}\n\n{_synth_prompt}",
                                    })
                                    _synth_result = await provider.chat(
                                        messages=[ChatMessage(**m) for m in _synth_msgs],
                                        model=model,
                                        max_tokens=max_output,
                                        temperature=request.temperature or 0.7,
                                        system_prompt=system_prompt,
                                    )
                                    cascade_costs.append(_synth_result.get("cost_usd", 0))
                                    cascade_tokens.append(
                                        _synth_result["usage"].get("total_tokens", 0))
                                    _synth_content = _synth_result.get("content", "").strip()
                                    if _synth_content and len(_synth_content) > 50:
                                        log.info(f"[{request_id}] Direct JSON→search→synthesis: "
                                                 f"{len(_synth_content)} chars")
                                        response_content = _synth_content
                                        llm_result = _synth_result
                                        tool_calls = None  # Skip normal tool processing
                                        # We'll handle source footer below
                                        _direct_synthesis_done = True
                        except (json.JSONDecodeError, TypeError, Exception) as e:
                            log.warning(f"[{request_id}] Direct JSON extraction failed: {e}")
                            _direct_synthesis_done = False
                    else:
                        _direct_synthesis_done = False

                    if not _direct_synthesis_done:
                        # Retry: add explicit "no code/JSON" instruction and resend
                        _retry_msgs = list(trimmed_messages)
                        _retry_msgs.append(ChatMessage(
                            role="user",
                            content=(
                                "WICHTIG: Antworte als normaler Text. "
                                "Gib KEINEN Python-Code, KEINE API-Aufrufe, KEIN print() aus. "
                                "Gib KEIN rohes JSON aus — JSON ist nur für tool_calls, nicht für Textantworten. "
                                "Wenn du eine Websuche brauchst, nutze das web_search Tool direkt als function call. "
                                "Beantworte die Frage direkt."
                            ),
                        ))
                        _retry_result = await provider.chat(
                            messages=_retry_msgs,
                            model=model,
                            max_tokens=max_output,
                            temperature=request.temperature or 0.7,
                            system_prompt=system_prompt,
                            tools=TOOL_DEFINITIONS if use_tools else None,
                        )
                        cascade_costs.append(_retry_result.get("cost_usd", 0))
                        cascade_tokens.append(
                            _retry_result["usage"].get("total_tokens", 0))

                        # Check if retry produced tool calls or real text
                        _retry_tools = _retry_result.get("tool_calls")
                        _retry_content = _retry_result.get("content", "").strip()

                        if _retry_tools:
                            # Model now correctly uses tools — process them
                            log.info(f"[{request_id}] Hallucination retry → "
                                     f"model now calls tools: "
                                     f"{[tc.get('function',{}).get('name','?') for tc in _retry_tools]}")
                            llm_result = _retry_result
                            tool_calls = _retry_tools
                            response_content = _retry_content
                            # Re-enter tool execution path
                            tool_results = await execute_tool_calls(tool_calls)
                            tools_used = [tr["name"] for tr in tool_results]

                            # Build follow-up for synthesis
                            follow_up = []
                            _search_depth = "snippets"
                            _analysis_mode = "factual"
                            for tc in tool_calls:
                                if tc.get("function", {}).get("name") == "web_search":
                                    try:
                                        _tc_args = json.loads(
                                            tc.get("function", {}).get("arguments", "{}"))
                                        _search_depth = _tc_args.get("search_depth",
                                                        _tc_args.get("depth", "snippets"))
                                        _depth_map = {"min": "snippets", "medium": "deep", "max": "thorough"}
                                        _search_depth = _depth_map.get(_search_depth, _search_depth)
                                        _analysis_mode = _tc_args.get("analysis_mode", "factual")
                                    except (json.JSONDecodeError, AttributeError):
                                        pass

                            _output_arch = extract_output_architecture(response_content) if response_content else {}
                            _blueprint = extract_blueprint(response_content) if (response_content and not _output_arch) else ""
                            _ctx_strategy = extract_context_strategy(response_content) if response_content else {"mode": "full"}

                            for m in trimmed_messages:
                                follow_up.append(_msg_to_dict(m))
                            follow_up.append({
                                "role": "assistant",
                                "content": response_content or "Ich suche nach Informationen...",
                                "tool_calls": [
                                    {"id": tc.get("id", ""), "type": "function",
                                     "function": tc.get("function", {})}
                                    for tc in tool_calls
                                ],
                            })
                            for tr in tool_results:
                                log.info(f"[{request_id}] Tool result [{tr['name']}]: "
                                         f"{tr['result'][:200]}...")
                                follow_up.append({
                                    "role": "tool",
                                    "tool_call_id": tr["tool_call_id"],
                                    "content": tr["result"],
                                })

                            synthesis_prompt = build_synthesis_prompt(
                                _search_depth, _analysis_mode, _blueprint,
                                output_architecture=_output_arch or None)
                            if synthesis_prompt:
                                follow_up.append({
                                    "role": "user", "content": synthesis_prompt,
                                })

                            await asyncio.sleep(1.0)
                            llm_result = await provider.chat(
                                messages=[], model=model,
                                max_tokens=max_output,
                                temperature=request.temperature or 0.7,
                                raw_messages=follow_up,
                            )
                            cascade_costs.append(llm_result.get("cost_usd", 0))
                            cascade_tokens.append(
                                llm_result["usage"].get("total_tokens", 0))
                            log.info(f"[{request_id}] Hallucination retry → "
                                     f"tool call complete: {tools_used}")

                        elif not any(p.lower() in _retry_content.lower()
                                     for p in _hallucination_patterns):
                            # Also verify no JSON hallucination in retry
                            _retry_is_json = (
                                _retry_content.strip().startswith("{")
                                and _retry_content.strip().endswith("}")
                                and any(k in _retry_content for k in
                                        ['"query"', '"output_architecture"', '"search_depth"'])
                            )
                            if _retry_is_json:
                                log.warning(f"[{request_id}] Hallucination retry also "
                                            f"produced JSON → using original")
                            else:
                                # Retry produced a real text answer
                                log.info(f"[{request_id}] Hallucination retry → "
                                         f"real text answer ({len(_retry_content)} chars)")
                                llm_result = _retry_result
                                response_content = _retry_content
                        else:
                            log.warning(f"[{request_id}] Hallucination retry also "
                                        f"produced code → using original")

            # ── Guard: Model refused to search and asked for clarification ──
            # Sometimes the model says "I can search for you" or "give me more details"
            # instead of just using the web_search tool. Auto-trigger search in this case.
            # SKIP for agent sessions (OpenClaw) — clarification questions are normal.
            if (not tool_calls and not escalated_to and use_tools
                    and response_content and len(response_content) < 600
                    and not _is_agent_session):
                _refusal_patterns = [
                    # German
                    "kann ich eine", "kann eine suche", "web-suche durchführen",
                    "websuche durchführen", "suche durchführen",
                    "benötige ich weitere", "benötige weitere",
                    "geben sie mir", "gib mir mehr details",
                    "können sie mir", "kannst du mir mehr",
                    "welche art von", "welche spezifischen",
                    "weitere informationen benötigt", "bitte teilen sie mir",
                    "um ihnen besser helfen", "um dir besser helfen",
                    # English (model sometimes responds in English to German questions)
                    "i can search", "i can perform", "could you provide more",
                    "need more information", "need more details",
                    "please provide more", "can you clarify",
                    "i would need more", "could you specify",
                    "what specific", "which specific",
                    "to help you better", "to better assist",
                    "i'll need to know", "i need to know more",
                    "can you tell me more", "what kind of",
                    "what type of", "do you have any specific",
                    "could you elaborate", "more context would help",
                    "i'd be happy to search", "i can look",
                    "shall i search", "would you like me to search",
                    "let me know if you", "let me know what",
                ]
                _rc_lower = response_content.lower()
                _is_refusal = any(p in _rc_lower for p in _refusal_patterns)
                if _is_refusal:
                    log.warning(f"[{request_id}] Model refused to search → "
                                f"auto-triggering web_search for '{user_query[:60]}'")
                    metrics.increment("refused_to_search_retries")
                    try:
                        _ws_result = await tool_web_search(query=user_query, depth="deep")
                        if _ws_result and len(_ws_result) > 100:
                            _synth_prompt = build_synthesis_prompt("deep")
                            _auto_msgs = []
                            for m in trimmed_messages:
                                _auto_msgs.append({"role": m.role, "content": m.text_content})
                            _auto_msgs.append({
                                "role": "user",
                                "content": f"{_ws_result}\n\n{_synth_prompt}",
                            })
                            _auto_result = await provider.chat(
                                messages=[ChatMessage(**m) for m in _auto_msgs],
                                model=model,
                                max_tokens=max_output,
                                temperature=request.temperature or 0.7,
                                system_prompt=system_prompt,
                            )
                            cascade_costs.append(_auto_result.get("cost_usd", 0))
                            cascade_tokens.append(
                                _auto_result["usage"].get("total_tokens", 0))
                            _auto_text = _auto_result.get("content", "").strip()
                            if _auto_text and len(_auto_text) > 50:
                                llm_result = _auto_result
                                response_content = _auto_text
                                tools_used = ["web_search"]
                                log.info(f"[{request_id}] Auto-search succeeded: "
                                         f"{len(_auto_text)} chars")
                    except Exception as e:
                        log.warning(f"[{request_id}] Auto-search failed: {e}")

            if escalated_to:
                log.info(f"[{request_id}] Cascade: cheap → {escalated_to}")
                metrics.increment("cascade_escalations", tags={"to": escalated_to})
                tier = escalated_to

                # Rate limit & budget check for escalated tier
                rate_ok, rate_reason, _ = rate_limiter.check(tier, estimated_tokens_count)
                if not rate_ok:
                    raise HTTPException(429, f"Rate limit exceeded: {rate_reason}")

                esc_cost_est = _estimate_request_cost(tier, estimated_tokens_count)
                budget_result = budget_guard.check(tier, esc_cost_est)
                if not budget_result["allowed"]:
                    raise HTTPException(429, f"Budget limit: {budget_result['reason']}")

                # Re-trim context for escalated tier
                esc_system = STATIC_SYSTEM_PROMPT if tier == "premium" else ""
                trimmed_messages, max_output = context_budget.apply(
                    tier, request.messages, esc_system
                )

                esc_response_type = "code_suggestion" if tier == "premium" else response_type
                out_strategy = output_strategy.get_strategy(
                    esc_response_type, tier, request.file_context
                )
                max_output = min(max_output, out_strategy["max_output_tokens"])

                if out_strategy.get("inject_diff_instruction"):
                    if trimmed_messages:
                        last_msg = trimmed_messages[-1]
                        if last_msg.role == "user":
                            trimmed_messages[-1] = ChatMessage(
                                role="user",
                                content=DIFF_INSTRUCTION + "\n\n" + last_msg.text_content
                            )

                esc_provider = _get_provider_for_tier(tier)
                model = _get_model_for_tier(tier)

                # ── Web Enrichment for escalations needing current data ──
                # Pre-enrich with DDG search + trafilatura
                # (We do NOT use OpenRouter's $0.02 web plugin — gateway tools are free)
                web_enrichment_ctx = ""
                if (escalated_to in ("cheap_plus", "medium", "premium")
                        and not is_document_analysis
                        and _WEB_ENRICHMENT_AVAILABLE):
                    try:
                        enricher = get_web_enricher()
                        # Detect language from query
                        _lang = "de" if any(c in user_query.lower() for c in
                                           ["ä","ö","ü","ß","ich","und","der","die","das"]) else "en"
                        enrichment = await enricher.enrich_query(
                            user_query, deep=True, language=_lang,
                        )
                        if enrichment.has_data:
                            web_enrichment_ctx = enrichment.enriched_context
                            log.info(f"[{request_id}] Web enrichment: {enrichment.method} | "
                                     f"{len(enrichment.search_results)} sources, "
                                     f"{enrichment.sources_fetched} deep | "
                                     f"~{enrichment.token_estimate} tok | "
                                     f"{enrichment.total_time_ms:.0f}ms")
                            metrics.increment("web_enrichments",
                                              tags={"method": enrichment.method,
                                                    "tier": escalated_to})
                    except Exception as e:
                        log.warning(f"[{request_id}] Web enrichment failed: {e}")

                # Inject web context into last user message
                if web_enrichment_ctx and trimmed_messages:
                    last = trimmed_messages[-1]
                    if last.role == "user":
                        trimmed_messages[-1] = ChatMessage(
                            role="user",
                            content=last.text_content + "\n\n" + web_enrichment_ctx,
                        )

                llm_result = await esc_provider.chat(
                    messages=trimmed_messages,
                    model=model,
                    max_tokens=max_output,
                    temperature=request.temperature or 0.7,
                    system_prompt=esc_system,
                    use_cache=(tier == "premium"),
                )
                cascade_costs.append(llm_result.get("cost_usd", 0))
                cascade_tokens.append(llm_result["usage"].get("total_tokens", 0))

        # ─── Step 11c: Verification (v2: Draft & Verify) ──────────────
        # Only runs for cheap/medium responses with large context
        if enhanced_result and enhanced_result.strategy == "verify":
            should_regen, verify_cost, verify_tokens = await _run_verification(
                request_id, tier, user_query,
                llm_result.get("content", ""),
                context_map, enhanced_result, full_context_text,
                total_request_tokens,
            )
            cascade_costs.append(verify_cost)
            cascade_tokens.append(verify_tokens)

            if should_regen:
                log.info(f"[{request_id}] Verification failed → regenerating with premium")
                metrics.increment("verification_regenerations")
                tier = "premium"
                provider = _get_provider_for_tier(tier)
                model = _get_model_for_tier(tier)
                trimmed_messages, max_output = context_budget.apply(
                    tier, request_messages_for_llm, STATIC_SYSTEM_PROMPT
                )
                llm_result = await provider.chat(
                    messages=trimmed_messages,
                    model=model,
                    max_tokens=max_output,
                    temperature=request.temperature or 0.7,
                    system_prompt=STATIC_SYSTEM_PROMPT,
                    use_cache=True,
                )
                cascade_costs.append(llm_result.get("cost_usd", 0))
                cascade_tokens.append(llm_result["usage"].get("total_tokens", 0))

        # ─── Step 11d: Response Validation (free, local checks) ────────
        # Detect truncated/incomplete code from cheap/medium and escalate
        if tier in ("cheap", "cheap_plus", "medium"):
            # Clean common formatting artifacts (diff markers, etc.)
            _raw_content = llm_result.get("content", "")
            _cleaned_content, _was_cleaned = _clean_code_response(_raw_content)
            if _was_cleaned:
                log.info(f"[{request_id}] Cleaned code response "
                         f"({len(_raw_content)} → {len(_cleaned_content)} chars)")
                llm_result["content"] = _cleaned_content

            validation = validate_response(
                response_text=llm_result.get("content", ""),
                finish_reason=llm_result.get("finish_reason", "stop"),
                tier=tier,
            )

            # ── Code adequacy check: did we get code when we expected it? ──
            # If the prompt was routed as long_code but the response has no
            # code blocks or is very short, the model gave an explanation
            # instead of the actual program. Force a retry via stitch path.
            if _expects_long_code and not validation.should_escalate:
                _resp_text = llm_result.get("content", "")
                _has_code_blocks = "```" in _resp_text
                _resp_lines = _resp_text.count("\n") + 1
                if not _has_code_blocks or _resp_lines < 30:
                    log.warning(
                        f"[{request_id}] Code adequacy FAIL: expected long code "
                        f"but got {len(_resp_text)} chars, {_resp_lines} lines, "
                        f"has_code_blocks={_has_code_blocks} → forcing retry")
                    # Override validation to trigger the stitch/repair cascade
                    validation.should_escalate = True
                    validation.is_valid = False
                    validation.reason = (
                        f"code_adequacy: expected code but got "
                        f"{_resp_lines} lines, has_code={_has_code_blocks}"
                    )
                    validation.details = [validation.reason]

            if validation.should_escalate:
                _resp_content = llm_result.get("content", "")
                _resp_finish = llm_result.get("finish_reason", "stop")

                log.info(f"[{request_id}] Validation failed on {tier}: "
                         f"finish={_resp_finish}, len={len(_resp_content)}, "
                         f"details={validation.details[:3]}")

                # ── Code adequacy retry: model gave text instead of code ──
                # If the prompt expected long code but got no code blocks,
                # retry once on the same tier before trying stitch/repair/premium.
                # A retry at $0.001 is much cheaper than premium at $0.10+.
                if _expects_long_code and "code_adequacy" in validation.reason:
                    log.info(f"[{request_id}] Code adequacy retry: "
                             f"resending to {tier} with reinforced prompt")
                    metrics.increment("code_adequacy_retries",
                                      tags={"tier": tier})
                    # Reinforce the prompt: append an explicit code instruction
                    _retry_msgs = list(trimmed_messages)
                    _retry_msgs.append({
                        "role": "user",
                        "content": (
                            "Bitte antworte NUR mit dem vollständigen Code als "
                            "```python Code-Block. Keine Erklärungen, "
                            "kein Text — nur der Code."
                        ),
                    })
                    _retry_result = await provider.chat(
                        messages=_retry_msgs,
                        model=model,
                        max_tokens=max_output,
                        temperature=request.temperature or 0.7,
                        system_prompt=system_prompt,
                    )
                    cascade_costs.append(_retry_result.get("cost_usd", 0))
                    cascade_tokens.append(
                        _retry_result["usage"].get("total_tokens", 0))

                    _retry_text = _retry_result.get("content", "")
                    _retry_has_code = "```" in _retry_text
                    _retry_lines = _retry_text.count("\n") + 1
                    if _retry_has_code and _retry_lines >= 30:
                        log.info(f"[{request_id}] Code adequacy retry SUCCESS: "
                                 f"{_retry_lines} lines, {len(_retry_text)} chars")
                        llm_result = _retry_result
                        _resp_content = _retry_text
                        _resp_finish = _retry_result.get("finish_reason", "stop")
                        # Re-validate the retry result
                        validation = validate_response(
                            response_text=_retry_text,
                            finish_reason=_resp_finish,
                            tier=tier,
                        )
                        if not validation.should_escalate:
                            # Retry fixed it — skip stitch/repair/premium
                            metrics.increment("code_adequacy_retry_success",
                                              tags={"tier": tier})
                        else:
                            log.info(f"[{request_id}] Code adequacy retry "
                                     f"produced code but validation fails: "
                                     f"{validation.reason} → continuing cascade")
                    else:
                        log.warning(f"[{request_id}] Code adequacy retry FAILED: "
                                    f"{_retry_lines} lines, has_code={_retry_has_code} "
                                    f"→ continuing to stitch/repair/premium")

                # ── Try code stitching BEFORE escalating to premium ──
                # Attempt stitching when:
                #   1. finish_reason=length AND code present (hard truncation), OR
                #   2. finish_reason=stop BUT validator says code is incomplete
                #      (model stopped voluntarily before finishing)
                # In both cases, calling cheap model again is cheaper than premium.
                _is_hard_truncation = _is_code_truncation(_resp_content, _resp_finish)
                _is_incomplete_code = _has_code_content(_resp_content) and not _is_hard_truncation

                # Init stitch state (used by repair step below)
                stitch_ok = False
                stitched = _resp_content
                stitch_validation = None

                if _is_hard_truncation or _is_incomplete_code:
                    _reason = "hard_truncation" if _is_hard_truncation else "incomplete_code"
                    log.info(f"[{request_id}] Code stitch candidate: "
                             f"reason={_reason}, len={len(_resp_content)}, "
                             f"finish={_resp_finish} "
                             f"→ trying stitch before premium escalation")
                    metrics.increment("code_stitch_attempts",
                                      tags={"tier": tier, "reason": _reason})

                    stitch_provider = _get_provider_for_tier(tier)
                    stitch_model = _get_model_for_tier(tier)
                    stitched, stitch_costs, stitch_tokens, stitch_ok = \
                        await _try_code_stitching(
                            request_id=request_id,
                            provider=stitch_provider,
                            model=stitch_model,
                            messages=trimmed_messages,
                            system_prompt=system_prompt,
                            partial_content=_resp_content,
                            max_tokens=max_output,
                            temperature=request.temperature or 0.7,
                            reason=_reason,
                        )
                    cascade_costs.extend(stitch_costs)
                    cascade_tokens.extend(stitch_tokens)

                    if stitch_ok:
                        # Re-validate the stitched result
                        stitch_validation = validate_response(
                            response_text=stitched,
                            finish_reason="stop",  # Stitching completed
                            tier=tier,
                        )
                        if not stitch_validation.should_escalate:
                            # Stitching fixed it — update result, skip premium
                            llm_result["content"] = stitched
                            llm_result["finish_reason"] = "stop"
                            log.info(f"[{request_id}] Code stitch SUCCESS: "
                                     f"{len(_resp_content)} → {len(stitched)} chars, "
                                     f"stitch_cost=${sum(stitch_costs):.6f} "
                                     f"(saved premium escalation)")
                            metrics.increment("code_stitch_success",
                                              tags={"tier": tier})
                            code_stitched = True
                            validation = stitch_validation  # Override → no escalation
                        else:
                            log.warning(f"[{request_id}] Code stitch completed but "
                                        f"validation still fails: "
                                        f"{stitch_validation.reason} → escalating")
                            metrics.increment("code_stitch_insufficient",
                                              tags={"tier": tier})
                    else:
                        log.warning(f"[{request_id}] Code stitch failed → "
                                    f"trying repair before premium")
                        metrics.increment("code_stitch_failed",
                                          tags={"tier": tier})

                # ── Try code REPAIR if stitch didn't fix validation ──
                # Send the broken code + validator errors back to cheap model
                # to fix (cheaper than premium re-generation)
                if validation.should_escalate and _has_code_content(
                        llm_result.get("content", "")):
                    _repair_content = (stitched if stitch_ok
                                       else llm_result.get("content", ""))
                    _repair_errors = (stitch_validation.details
                                      if stitch_ok and stitch_validation
                                      else validation.details)

                    if _repair_errors:
                        log.info(f"[{request_id}] Attempting code repair: "
                                 f"errors={_repair_errors[:3]}")
                        metrics.increment("code_repair_attempts",
                                          tags={"tier": tier})

                        repair_provider = _get_provider_for_tier(tier)
                        repair_model = _get_model_for_tier(tier)
                        repaired, repair_costs, repair_tokens, repair_ok = \
                            await _try_code_repair(
                                request_id=request_id,
                                provider=repair_provider,
                                model=repair_model,
                                messages=trimmed_messages,
                                system_prompt=system_prompt,
                                broken_content=_repair_content,
                                validation_errors=_repair_errors,
                                max_tokens=max_output,
                                temperature=request.temperature or 0.7,
                            )
                        cascade_costs.extend(repair_costs)
                        cascade_tokens.extend(repair_tokens)

                        if repair_ok:
                            repair_validation = validate_response(
                                response_text=repaired,
                                finish_reason="stop",
                                tier=tier,
                            )
                            if not repair_validation.should_escalate:
                                llm_result["content"] = repaired
                                llm_result["finish_reason"] = "stop"
                                log.info(f"[{request_id}] Code repair SUCCESS: "
                                         f"{len(_repair_content)} → {len(repaired)} chars, "
                                         f"repair_cost=${sum(repair_costs):.6f} "
                                         f"(saved premium escalation)")
                                metrics.increment("code_repair_success",
                                                  tags={"tier": tier})
                                code_repaired = True
                                validation = repair_validation
                            else:
                                log.warning(f"[{request_id}] Code repair completed but "
                                            f"validation still fails: "
                                            f"{repair_validation.reason} → premium")
                                metrics.increment("code_repair_insufficient",
                                                  tags={"tier": tier})
                        else:
                            log.warning(f"[{request_id}] Code repair failed → "
                                        f"escalating to premium")
                            metrics.increment("code_repair_failed",
                                              tags={"tier": tier})

            # If validation still requires escalation (stitch+repair didn't help or non-code)
            if validation.should_escalate:
                log.warning(f"[{request_id}] Response validation failed ({tier}): "
                            f"{validation.reason} → escalating to premium")
                metrics.increment("validation_escalations",
                                  tags={"from": tier, "reason": validation.details[0][:50] if validation.details else "unknown"})

                prev_tier = tier
                tier = "premium"
                validation_escalated = prev_tier  # Track for metadata
                provider = _get_provider_for_tier(tier)
                model = _get_model_for_tier(tier)

                trimmed_messages, max_output = context_budget.apply(
                    tier, request_messages_for_llm, STATIC_SYSTEM_PROMPT
                )

                llm_result = await provider.chat(
                    messages=trimmed_messages,
                    model=model,
                    max_tokens=max_output,
                    temperature=request.temperature or 0.7,
                    system_prompt=STATIC_SYSTEM_PROMPT,
                    use_cache=True,
                )
                cascade_costs.append(llm_result.get("cost_usd", 0))
                cascade_tokens.append(llm_result["usage"].get("total_tokens", 0))
                log.info(f"[{request_id}] Validation escalation: {prev_tier}→premium "
                         f"(+${llm_result.get('cost_usd', 0):.6f})")

        # ─── Step 11e: Post-Response Fact-Check (web verification) ────
        # For complex factual queries answered by medium/premium without
        # prior web enrichment, verify claims against web sources.
        # SKIP for: agent sessions (OpenClaw manages its own flow),
        #           API-sourced data (weather/stocks already factual),
        #           tool-call responses (synthesis already did web search)
        _skip_fact_check = (
            _is_agent_session
            or bool(tools_used)  # Already used tools (web_search, get_weather, etc.)
        )
        if (not _skip_fact_check
                and tier in ("medium", "premium")
                and not web_enrichment_ctx  # No pre-enrichment was done
                and not is_document_analysis
                and _WEB_ENRICHMENT_AVAILABLE
                and response_type not in ("code_suggestion",)
                and len(llm_result.get("content", "")) > 200):
            classification = classify_needs_enrichment(user_query)
            if classification["needs_fact_check"]:
                try:
                    enricher = get_web_enricher()
                    _lang = "de" if any(c in user_query.lower() for c in
                                       ["ä","ö","ü","ß","ich","und","der","die","das"]) else "en"
                    fc_result = await enricher.fact_check(
                        user_query, llm_result["content"], language=_lang,
                    )
                    if fc_result and fc_result.has_data:
                        log.info(f"[{request_id}] Fact-check: {fc_result.sources_fetched} "
                                 f"deep sources, ~{fc_result.token_estimate} tok | "
                                 f"re-calling {tier} with verification context")
                        metrics.increment("fact_checks", tags={"tier": tier})

                        # Re-call LLM with original response + fact-check context
                        fc_messages = list(trimmed_messages) if trimmed_messages else []
                        # Add the LLM's initial response
                        fc_messages.append(ChatMessage(
                            role="assistant",
                            content=llm_result["content"],
                        ))
                        # Add fact-check context as follow-up user message
                        fc_messages.append(ChatMessage(
                            role="user",
                            content=(
                                fc_result.enriched_context + "\n\n"
                                "Bitte überprüfe und korrigiere deine vorherige Antwort "
                                "anhand dieser Quellen. Wenn Korrekturen nötig sind, "
                                "gib die vollständig korrigierte Antwort. Wenn alles stimmt, "
                                "wiederhole die Antwort mit Quellenangaben."
                                if _lang == "de" else
                                "Please verify and correct your previous answer based on "
                                "these sources. If corrections are needed, provide the "
                                "fully corrected answer. If correct, repeat with citations."
                            ),
                        ))

                        fc_provider = _get_provider_for_tier(tier)
                        fc_model = _get_model_for_tier(tier)
                        fc_llm = await fc_provider.chat(
                            messages=fc_messages,
                            model=fc_model,
                            max_tokens=max_output,
                            temperature=0.3,  # Lower temp for factual accuracy
                            system_prompt=STATIC_SYSTEM_PROMPT if tier == "premium" else "",
                            use_cache=False,
                        )
                        cascade_costs.append(fc_llm.get("cost_usd", 0))
                        cascade_tokens.append(fc_llm["usage"].get("total_tokens", 0))
                        llm_result = fc_llm  # Use fact-checked version

                        log.info(f"[{request_id}] Fact-check complete: "
                                 f"+${fc_llm.get('cost_usd', 0):.6f}")
                except Exception as e:
                    log.warning(f"[{request_id}] Fact-check failed: {e}")

        # ─── Step 12: Record Costs & Cache ────────────────────────────
        actual_cost = sum(cascade_costs)
        actual_tokens = sum(cascade_tokens)

        rate_limiter.record(tier, actual_tokens, actual_cost)
        budget_guard.record_spend(actual_cost)
        metrics.increment("total_cost_usd", actual_cost)

        # Append source footer from web search (if enabled in config)
        # Always set _final_content from synthesis result
        _final_content = llm_result.get("content", "") or ""
        _filtered_footer = ""

        # ── Parse and strip quality self-assessment tag ──
        _quality_assessment = None
        if "<!--QUALITY:" in _final_content:
            import re as _re_q
            # Match <!--QUALITY:{...}--> with nested braces (retry object)
            _q_match = _re_q.search(
                r'<!--QUALITY:\s*(\{.+?\})\s*-->', _final_content, _re_q.DOTALL
            )
            if _q_match:
                _q_raw = _q_match.group(1)
                # Handle nested JSON: find balanced braces
                _depth = 0
                _end = 0
                for _ci, _ch in enumerate(_q_raw):
                    if _ch == '{': _depth += 1
                    elif _ch == '}': _depth -= 1
                    if _depth == 0:
                        _end = _ci + 1
                        break
                _q_json = _q_raw[:_end] if _end > 0 else _q_raw
                try:
                    _quality_assessment = json.loads(_q_json)
                    log.info(f"[{request_id}] Quality assessment: {_quality_assessment}")
                except (json.JSONDecodeError, TypeError):
                    log.warning(f"[{request_id}] Failed to parse quality tag: "
                                f"{_q_json[:150]}")
                # Strip the quality tag from user-visible content
                _final_content = _final_content[:_q_match.start()].rstrip()
                # Also strip any remaining quality tags
                _final_content = _re_q.sub(
                    r'<!--QUALITY:.*?-->', '', _final_content, flags=_re_q.DOTALL
                ).rstrip()

        # Log quality warnings for monitoring
        if _quality_assessment:
            _qa = _quality_assessment
            if not _qa.get("answered", True):
                log.warning(f"[{request_id}] Quality: NOT ANSWERED "
                            f"(confidence={_qa.get('confidence','?')}, "
                            f"sources={_qa.get('source_count',0)})")
            elif _qa.get("confidence") == "low":
                log.warning(f"[{request_id}] Quality: LOW CONFIDENCE "
                            f"(sources={_qa.get('source_count',0)}, "
                            f"has_facts={_qa.get('has_facts', False)})")

            # ── Quality-based retry: synthesis recommends follow-up search ──
            _retry = _qa.get("retry")
            if (_retry and isinstance(_retry, dict)
                    and _retry.get("queries")
                    and not getattr(main, '_quality_retry_done', False)):
                _retry_queries = _retry["queries"][:3]
                _retry_tf = _retry.get("time_filter", "w")
                _retry_reason = _retry.get("reason", "low quality")
                log.info(f"[{request_id}] Quality retry recommended: "
                         f"queries={_retry_queries} tf={_retry_tf} "
                         f"reason='{_retry_reason}'")

                try:
                    # Execute follow-up web search with recommended queries
                    import uuid as _uuid_r
                    _retry_tc_id = f"retry_{_uuid_r.uuid4().hex[:8]}"
                    _retry_tool_calls = [{
                        "id": _retry_tc_id,
                        "type": "function",
                        "function": {
                            "name": "web_search",
                            "arguments": json.dumps({
                                "query": _retry_queries[0],
                                "search_depth": "medium",
                                "time_filter": _retry_tf,
                                "additional_queries": _retry_queries[1:] if len(_retry_queries) > 1 else None,
                            }),
                        },
                    }]
                    _retry_results = await execute_tool_calls(_retry_tool_calls)
                    _retry_tool_text = _retry_results[0]["result"] if _retry_results else ""

                    if _retry_tool_text and len(_retry_tool_text) > 200:
                        # Re-synthesize with original + retry results
                        from datetime import datetime as _dt_r
                        _today_r = _dt_r.now().strftime("%d.%m.%Y")
                        _retry_synthesis_prompt = (
                            f"Du bist ein Recherche-Assistent. Heutiges Datum: {_today_r}.\n"
                            "Dir wurden die Ergebnisse einer NACHSUCHE bereitgestellt.\n"
                            "Du hast bereits eine erste Antwort gegeben (siehe unten), "
                            "aber die Quellen waren unzureichend.\n\n"
                            "Erstelle eine VERBESSERTE Antwort die die neuen Quellen einbezieht.\n"
                            "Behalte die Struktur bei: Überschriften, Fazit, [N] Quellen.\n"
                            "SPRACHE: Antworte in der Sprache der Nutzerfrage.\n"
                            "Halte die Antwort kompakt (Telegram-Kontext).\n\n"
                            "SELBST-BEWERTUNG: Füge am Ende ein:\n"
                            '<!--QUALITY:{"answered":true/false,"confidence":"high"/"medium"/"low",'
                            '"source_count":N,"has_facts":true/false,"retry":null}-->'
                        )
                        _retry_follow_up = [
                            {"role": "system", "content": _retry_synthesis_prompt},
                            {"role": "user", "content": (
                                f"Ursprüngliche Frage: {_routing_query}\n\n"
                                f"Erste Antwort (unzureichend):\n{_final_content}\n\n"
                                f"Neue Suchergebnisse:\n{_retry_tool_text}"
                            )},
                        ]
                        await asyncio.sleep(0.5)
                        _retry_llm = await provider.chat(
                            messages=[],
                            model=model,
                            max_tokens=max_output,
                            temperature=0.7,
                            raw_messages=_retry_follow_up,
                        )
                        cascade_costs.append(_retry_llm.get("cost_usd", 0))
                        cascade_tokens.append(
                            _retry_llm["usage"].get("total_tokens", 0))

                        _retry_content = _retry_llm.get("content", "") or ""
                        # Parse quality tag from retry
                        import re as _re_qr
                        _qr_match = _re_qr.search(
                            r'<!--QUALITY:\s*(\{.+?\})\s*-->',
                            _retry_content, _re_qr.DOTALL
                        )
                        _retry_quality = None
                        if _qr_match:
                            _qr_raw = _qr_match.group(1)
                            _d = 0
                            _e = 0
                            for _j, _c in enumerate(_qr_raw):
                                if _c == '{': _d += 1
                                elif _c == '}': _d -= 1
                                if _d == 0: _e = _j + 1; break
                            try:
                                _retry_quality = json.loads(
                                    _qr_raw[:_e] if _e > 0 else _qr_raw)
                            except (json.JSONDecodeError, TypeError):
                                pass
                            _retry_content = _retry_content[:_qr_match.start()].rstrip()
                            _retry_content = _re_qr.sub(
                                r'<!--QUALITY:.*?-->', '',
                                _retry_content, flags=_re_qr.DOTALL
                            ).rstrip()

                        # Use retry result if it's better
                        _retry_better = (
                            _retry_quality
                            and _retry_quality.get("answered", False)
                            and _retry_quality.get("confidence", "low") != "low"
                        )
                        if _retry_better or len(_retry_content) > len(_final_content) * 1.3:
                            _final_content = _retry_content
                            _quality_assessment = _retry_quality or _quality_assessment
                            # Update source footer
                            _source_footer_retry = get_last_source_footer()
                            if _source_footer_retry:
                                from tool_executor import get_last_source_footer as _glsf2
                                clear_source_footer()
                            log.info(f"[{request_id}] Quality retry SUCCESS: "
                                     f"used improved answer "
                                     f"(quality={_retry_quality})")
                        else:
                            log.info(f"[{request_id}] Quality retry: kept original "
                                     f"(retry not better, quality={_retry_quality})")

                    actual_cost = sum(cascade_costs)
                    actual_tokens = sum(cascade_tokens)
                except Exception as _retry_err:
                    log.warning(f"[{request_id}] Quality retry failed: {_retry_err}")
        try:
            _source_footer = get_last_source_footer()
            _has_web_tools = tools_used and ("web_search" in tools_used or "get_news" in tools_used)
            
            if _source_footer:
                log.info(f"[{request_id}] Source footer available: "
                         f"{len(_source_footer)} chars, {_source_footer.count('[')} refs | "
                         f"tools_used={tools_used} | agent_tc={bool(_agent_tool_calls)}")
            
            if _source_footer and _has_web_tools:
                import re as _re
                
                # Post-process: remove inline URL citations the model shouldn't have used
                _final_content = _re.sub(
                    r'\s*\(Quelle:\s*https?://[^)]+\)', '', _final_content)
                _final_content = _re.sub(
                    r'\s*\(Quelle:\s*[a-z0-9][a-z0-9.-]+\.[a-z]{2,}(?:\s*,\s*[a-z0-9.-]+\.[a-z]{2,})*\)',
                    '', _final_content, flags=_re.IGNORECASE)
                
                # ── Strip lazy-response trailing sentences ──
                # Model sometimes adds "Es ist ratsam, die Webseiten zu besuchen" etc.
                _lazy_patterns = [
                    r'\n+(?:Es ist ratsam|Bitte beachte|Es empfiehlt sich|Für aktuelle|'
                    r'Es lohnt sich|Weitere Informationen|Für weitere|Besuchen Sie|'
                    r'Für detaillierte)[^.]*\.\s*$',
                ]
                for lp in _lazy_patterns:
                    _final_content = _re.sub(lp, '', _final_content)
                
                # ── Strip model-generated source sections ──
                # The model sometimes generates its own "Quellen:" with hallucinated URLs.
                # We replace them with our real footer from build_source_footer().
                _source_section_patterns = [
                    r'\n---\s*\n\*?\*?Quellen:?\*?\*?\s*\n.*',  # Our format + bold
                    r'\nQuellen:\s*\n\[1\].*',                    # Standard
                    r'\nSources:\s*\n\[1\].*',                    # English
                    r'\nQuellen:\s*\n\*?\*?\[1\].*',              # Bold refs
                    r'\n\*\*Quellen:?\*\*\s*\n.*',                # Bold header
                    r'\n\*\*Sources:?\*\*\s*\n.*',                # Bold English
                    r'\nQuellen:\s*\n[•\-].*',                    # Bullet format
                    r'\nQuellen:\s*\n\d+\..*',                    # Numbered format
                    r'\nQuellen:\s*\nhttps?://.*',                # Direct URLs
                ]
                for pattern in _source_section_patterns:
                    _stripped = _re.sub(pattern, '', _final_content, flags=_re.DOTALL)
                    if len(_stripped) < len(_final_content):
                        log.info(f"[{request_id}] Stripped model-generated source section "
                                 f"({len(_final_content) - len(_stripped)} chars)")
                        _final_content = _stripped.rstrip()
                        break
                
                # ── Validate inline references [N] ──
                # Count real sources in the footer and strip refs that exceed the count.
                _real_source_count = _source_footer.count("\n[")
                if _real_source_count > 0:
                    # Find all [N] references in the response
                    _all_refs = set(int(m) for m in _re.findall(r'\[(\d+)\]', _final_content))
                    _invalid_refs = {n for n in _all_refs if n > _real_source_count or n < 1}
                    if _invalid_refs:
                        log.warning(f"[{request_id}] Invalid source refs {_invalid_refs} "
                                    f"(only {_real_source_count} real sources) → removing")
                        for ref_num in _invalid_refs:
                            _final_content = _final_content.replace(f"[{ref_num}]", "")
                        # Clean up double spaces left behind
                        _final_content = _re.sub(r'  +', ' ', _final_content)
                
                # Filter footer: only keep sources actually referenced in LLM response
                # Also renumbers both text and footer for consistency
                _final_content, _filtered_footer = _filter_source_footer(
                    _final_content, _source_footer
                )
                if _filtered_footer:
                    _final_content = _final_content.rstrip() + "\n" + _filtered_footer
                    log.info(f"[{request_id}] Source footer appended "
                             f"({_filtered_footer.count('[')}) sources, "
                             f"filtered from {_source_footer.count('[')}")
            clear_source_footer()
        except Exception:
            pass  # Source footer is optional, never fail on it

        # Store in caches (skip if no_cache mode or agent session)
        response_data = {"content": _final_content}

        if config.cache.exact_cache_enabled and not no_cache and not _is_agent_session:
            exact_cache.set(
                user_query, fingerprint, response_data,
                response_type, model
            )

        if config.cache.semantic_cache_enabled and not no_cache and not _is_agent_session and tier in ("cheap", "cheap_plus", "medium", "premium"):
            embedding_fn = None
            if "embedding" in providers:
                embedding_fn = providers["embedding"].get_embedding
            await semantic_cache.set(
                user_query, fingerprint, response_data,
                response_type, model, embedding_fn
            )

        # Store for idempotency
        idempotency.store(idem_key, response_data)

        latency = (time.time() - start_time) * 1000
        metrics.histogram("request_latency_ms", latency, tags={"tier": tier})

        cascade_info = ""
        if cascade_mode:
            if size_escalated_from:
                cascade_info = f" | size→{tier} ({total_request_tokens}tok)"
            elif escalated_to:
                cascade_info = f" | cascade→{escalated_to}"
            else:
                cascade_info = " | cascade:direct"
        log.info(f"[{request_id}] {tier}/{model} | {actual_tokens} tok | ${actual_cost:.4f} | {latency:.0f}ms{cascade_info}"
                 f"{' | web_enriched' if web_enrichment_ctx else ''}"
                 f"{' | media:' + ','.join(media_types) if has_media else ''}")

        # ── Debug buffer capture ──
        try:
            _debug_entry = {
                "id": request_id,
                "ts": datetime.utcnow().isoformat() + "Z",
                "query": user_query[:300],
                "tier": tier,
                "model": model,
                "cost_usd": actual_cost,
                "latency_ms": round(latency, 1),
                "tokens": {"in": llm_result["usage"]["prompt_tokens"],
                           "out": llm_result["usage"]["completion_tokens"],
                           "reasoning": llm_result["usage"].get("reasoning_tokens", 0)},
                "escalated": escalated_to,
                "tools_used": tools_used,
                "is_agent": _is_agent_session,
                "system_prompt": (system_prompt or "")[:3000],
                "messages": [
                    {"role": m.role,
                     "content": (m.text_content or "")[:800],
                     "has_tc": bool(getattr(m, 'tool_calls', None)),
                     "tc_names": [tc.get("function", {}).get("name", "?")
                                  for tc in (getattr(m, 'tool_calls', None) or [])],
                     "has_tid": bool(getattr(m, 'tool_call_id', None))}
                    for m in (trimmed_messages or [])
                ],
                "response": (_final_content or "")[:3000],
                "response_tool_calls": (
                    [tc.get("function", {}).get("name", "?")
                     for tc in (_agent_tool_calls or llm_result.get("tool_calls") or [])]
                    if (_agent_tool_calls or llm_result.get("tool_calls")) else None
                ),
                "routing": {
                    "action": (enhanced_result.action.value if enhanced_result
                              else route_result.action.value if route_result else None),
                    "response_type": (enhanced_result.response_type if enhanced_result
                                     else route_result.response_type if route_result
                                     else response_type),
                    "is_code": (enhanced_result.is_code_generation if enhanced_result
                               else route_result.is_code_generation if route_result
                               and hasattr(route_result, 'is_code_generation') else False),
                    "needs_web": _needs_web_search if cascade_mode else None,
                },
                "web_enrichment": web_enrichment_ctx[:200] if web_enrichment_ctx else None,
                "quality": _quality_assessment,
            }
            _debug_buffer.append(_debug_entry)
        except Exception as _de:
            log.debug(f"[{request_id}] Debug capture failed: {_de}")

        # ── Agent pass-through: return tool_calls directly to client ──
        if _agent_tool_calls:
            from starlette.responses import JSONResponse as _JSONResponse
            _agent_msg = {"role": "assistant", "content": _final_content or None}
            _agent_msg["tool_calls"] = _agent_tool_calls
            
            if request.stream:
                import json as _json
                _created = int(time.time())
                
                async def _agent_stream():
                    if _final_content:
                        yield f"data: {_json.dumps({'id': request_id, 'object': 'chat.completion.chunk', 'created': _created, 'model': model, 'choices': [{'index': 0, 'delta': {'role': 'assistant', 'content': _final_content}, 'finish_reason': None}]})}\n\n"
                    yield f"data: {_json.dumps({'id': request_id, 'object': 'chat.completion.chunk', 'created': _created, 'model': model, 'choices': [{'index': 0, 'delta': {'tool_calls': _agent_tool_calls}, 'finish_reason': 'tool_calls'}]})}\n\n"
                    yield "data: [DONE]\n\n"
                
                from starlette.responses import StreamingResponse as _StreamResp
                log.info(f"[{request_id}] Agent pass-through stream with "
                         f"{len(_agent_tool_calls)} tool_calls")
                return _StreamResp(_agent_stream(), media_type="text/event-stream")
            
            _agent_response = {
                "id": request_id,
                "object": "chat.completion",
                "created": int(time.time()),
                "model": model,
                "choices": [{
                    "index": 0,
                    "message": _agent_msg,
                    "finish_reason": "tool_calls",
                }],
                "usage": {
                    "prompt_tokens": llm_result["usage"]["prompt_tokens"],
                    "completion_tokens": llm_result["usage"]["completion_tokens"],
                    "total_tokens": actual_tokens,
                },
            }
            log.info(f"[{request_id}] Agent pass-through response with "
                     f"{len(_agent_tool_calls)} tool_calls")
            return _JSONResponse(content=_agent_response)

        return _build_response(
            request_id, response_data, model,
            {
                "tier": tier,
                "cost_usd": actual_cost,
                "latency_ms": round(latency, 1),
                "cached": False,
                "cache_read_tokens": llm_result["usage"].get("cache_read_tokens", 0),
                "output_strategy": out_strategy["mode"],
                "cascade_escalated": escalated_to,
                "cascade_cheap_cost": round(cascade_costs[0], 6) if escalated_to else None,
                "size_escalated": size_escalated_from is not None,
                "multimodal": has_media or (vision_result is not None),
                "media_types": media_types if media_types else None,
                "vision_strategy": vision_result["strategy"] if vision_result else None,
                "vision_image_skipped": vision_result["can_skip_image"] if vision_result else None,
                "tool_calls_executed": tools_used if tools_used else None,
                "validation_escalated": validation_escalated,  # e.g. "medium" if code was incomplete
                "code_stitched": code_stitched,  # True if truncated code was continued instead of escalated
                "code_repaired": code_repaired,  # True if cheap model fixed its own broken code
                "web_enrichment": web_enrichment_ctx[:50] + "..." if web_enrichment_ctx else None,
                "synthesis_mode": f"{_search_depth}/{_analysis_mode}/ctx={_ctx_mode}" if tools_used and "web_search" in tools_used else None,
                "request_tokens": total_request_tokens,
                # Routing classification info (for test console debugging)
                "routing": {
                    "action": (enhanced_result.action.value if enhanced_result
                              else route_result.action.value if route_result else None),
                    "response_type": (enhanced_result.response_type if enhanced_result
                                     else route_result.response_type if route_result else response_type),
                    "is_code": (enhanced_result.is_code_generation if enhanced_result
                               else route_result.is_code_generation if route_result
                               and hasattr(route_result, 'is_code_generation') else False),
                    "layer": (enhanced_result.routing_layer if enhanced_result
                              else "cascade" if not route_result else "intent_router"),
                },
                # v2: Enhanced routing info
                "enhanced_routing": {
                    "layer": enhanced_result.routing_layer if enhanced_result else None,
                    "strategy": enhanced_result.strategy if enhanced_result else None,
                    "utility_score": enhanced_result.utility_score if enhanced_result else None,
                    "chunks_retrieved": len(enhanced_result.needed_chunks) if enhanced_result else 0,
                    "context_map_type": context_map.document_type if context_map else None,
                    "context_map_chunks": context_map.chunk_count if context_map else 0,
                } if enhanced_result or context_map else None,
            },
            usage=UsageInfo(
                prompt_tokens=llm_result["usage"]["prompt_tokens"],
                completion_tokens=llm_result["usage"]["completion_tokens"],
                total_tokens=actual_tokens,
                cache_read_tokens=llm_result["usage"].get("cache_read_tokens", 0),
                cache_write_tokens=llm_result["usage"].get("cache_write_tokens", 0),
                reasoning_tokens=llm_result["usage"].get("reasoning_tokens", 0),
                estimated_cost_usd=actual_cost,
            ),
            stream=request.stream,
        )

    except HTTPException:
        raise
    except ProviderRateLimitError as e:
        metrics.increment("provider_rate_limits", tags={"provider": e.provider, "model": e.model})
        log.warning(f"[{request_id}] Provider rate limited: {e}")
        retry_after = int(e.retry_after) if e.retry_after else 10
        raise HTTPException(
            status_code=429,
            detail=f"Provider rate limited ({e.provider}/{e.model}). "
                   f"Retry after {retry_after}s.",
            headers={"Retry-After": str(retry_after)},
        )
    except ProviderOverloadError as e:
        metrics.increment("provider_overloads", tags={"provider": e.provider})
        log.warning(f"[{request_id}] Provider overloaded: {e}")
        raise HTTPException(
            status_code=503,
            detail=f"Provider temporarily overloaded ({e.provider}). Please retry.",
            headers={"Retry-After": "5"},
        )
    except httpx.HTTPStatusError as e:
        # Catch any other HTTP errors from providers
        metrics.increment("errors")
        status = e.response.status_code
        log.error(f"[{request_id}] Provider HTTP error {status}: {e}")
        # Don't leak provider URLs to client
        if status in (429, 503):
            raise HTTPException(status_code=status,
                                detail=f"Provider temporarily unavailable ({status}). Please retry.",
                                headers={"Retry-After": "10"})
        if status == 500:
            # Provider internal error — often transient, suggest retry
            raise HTTPException(502,
                                f"Provider returned error 500 (transient). Please retry.",
                                headers={"Retry-After": "5"})
        raise HTTPException(502, f"Provider returned error {status}")
    except httpx.ConnectError as e:
        metrics.increment("errors")
        log.error(f"[{request_id}] Provider connection failed: {e}")
        raise HTTPException(502, "Cannot reach LLM provider")
    except httpx.TimeoutException as e:
        metrics.increment("errors")
        log.error(f"[{request_id}] Provider timeout: {e}")
        raise HTTPException(504, "LLM provider timeout")
    except Exception as e:
        metrics.increment("errors")
        import traceback
        tb = traceback.format_exc()
        log.error(f"[{request_id}] Error: {e}\n{tb}")
        # Include enough detail to debug without leaking sensitive info
        error_type = type(e).__name__
        error_detail = str(e)[:150]
        # Find the relevant line in traceback
        tb_lines = tb.strip().split('\n')
        last_code_line = ""
        for line in reversed(tb_lines):
            ls = line.strip()
            if ls.startswith("File") and any(f in ls for f in
                ["main.py", "providers.py", "enhanced_router.py", "context.py",
                 "context_mapper.py", "router.py", "rate_limiter.py"]):
                last_code_line = ls
                break
        detail = f"Internal gateway error: {error_type}: {error_detail}"
        if last_code_line:
            detail += f" | at: {last_code_line}"
        raise HTTPException(500, detail)
    finally:
        metrics.gauge("active_requests", 0)


# ─── Helper Functions ─────────────────────────────────────────────────────

async def _run_verification(
    request_id: str, tier: str, user_query: str, draft_content: str,
    context_map, enhanced_result, full_context_text: str, context_tokens: int,
) -> tuple[bool, float, int]:
    """Run Draft & Verify. Returns (should_regenerate, cost, tokens)."""
    if not verification_layer.should_verify(tier, context_tokens):
        return False, 0.0, 0

    # Build targeted context for verification
    verify_context = ""
    if enhanced_result and enhanced_result.needed_chunks and context_map:
        verify_context = context_mapper.retrieve_chunks(
            full_context_text, context_map, enhanced_result.needed_chunks, margin_chars=300
        )
    elif full_context_text:
        max_chars = 6000
        if len(full_context_text) > max_chars:
            half = max_chars // 2
            verify_context = full_context_text[:half] + "\n...\n" + full_context_text[-half:]
        else:
            verify_context = full_context_text

    if not verify_context:
        return False, 0.0, 0

    # Verifier is one tier above the drafter
    verifier_role = "medium" if tier in ("cheap", "cheap_plus") else "premium" if tier == "medium" else None
    if not verifier_role:
        return False, 0.0, 0

    verifier_provider = providers.get(verifier_role)
    verifier_model = _get_model_for_tier(verifier_role)
    if not verifier_provider:
        return False, 0.0, 0

    log.info(f"[{request_id}] Verifying draft with {verifier_role}/{verifier_model}")
    v_result = await verification_layer.verify_draft(
        query=user_query, draft_response=draft_content,
        context_text=verify_context,
        verifier_provider=verifier_provider, verifier_model=verifier_model,
    )
    log.info(f"[{request_id}] Verification: {v_result.verdict} (conf={v_result.confidence:.2f}, regen={v_result.should_regenerate})")
    return v_result.should_regenerate, v_result.verification_cost, v_result.verification_tokens

def _get_provider_for_tier(tier: str) -> LLMProvider:
    """Get the LLM provider for a given tier."""
    tier_to_role = {
        "local": "cheap",
        "cheap": "cheap",
        "cheap_plus": "cheap_plus",
        "medium": "medium",
        "premium": "premium",
    }
    role = tier_to_role.get(tier, "cheap")
    return providers.get(role, providers.get("cheap"))


def _get_model_for_tier(tier: str) -> str:
    """Get the model name for a given tier."""
    tier_to_model = {
        "local": config.providers.cheap_model,
        "cheap": config.providers.cheap_model,
        "cheap_plus": config.providers.cheap_plus_model,
        "medium": config.providers.medium_model,
        "premium": config.providers.premium_model,
    }
    return tier_to_model.get(tier, config.providers.cheap_model)


def _estimate_request_cost(tier: str, tokens: int) -> float:
    """Rough cost estimate for budget check."""
    prices = {
        "local": 0,
        "cheap": 0.10 / 1_000_000,
        "cheap_plus": 0.15 / 1_000_000,  # Gemini 2 Flash ($0.10/$0.40/M) + web search overhead
        "medium": 1.0 / 1_000_000,  # Gemini 3 Flash ($0.50/$3.00/M blended)
        "premium": 9.0 / 1_000_000,  # Blended input+output
    }
    return tokens * prices.get(tier, 0)


def _build_response(request_id: str, response_data: dict, model: str,
                    metadata: dict = None, usage: UsageInfo = None,
                    stream: bool = False):
    """Build an OpenAI-compatible response (JSON or SSE stream)."""
    content = response_data.get("content", "")

    if stream:
        return _to_sse_stream(request_id, content, model, metadata, usage)

    response = ChatResponse(
        id=request_id,
        created=int(time.time()),
        model=model,
        choices=[ChatChoice(
            index=0,
            message=ChatMessage(role="assistant", content=content),
            finish_reason="stop",
        )],
        usage=usage or UsageInfo(),
        gateway_metadata=metadata,
    )
    # Use exclude_none to avoid sending tool_calls=null, tool_call_id=null etc.
    # which breaks OpenAI-compatible clients (OpenClaw)
    from starlette.responses import JSONResponse as _JSONResp
    return _JSONResp(content=response.model_dump(exclude_none=True))


def _to_sse_stream(request_id: str, content: str, model: str,
                   metadata: dict = None, usage: UsageInfo = None) -> StreamingResponse:
    """Convert a completed response to SSE stream format for OpenAI compatibility."""
    import json as _json
    created = int(time.time())

    async def generate():
        # Single content chunk (gateway doesn't do real streaming from upstream)
        chunk = {
            "id": request_id,
            "object": "chat.completion.chunk",
            "created": created,
            "model": model,
            "choices": [{
                "index": 0,
                "delta": {"role": "assistant", "content": content},
                "finish_reason": None,
            }],
        }
        yield f"data: {_json.dumps(chunk)}\n\n"

        # Final chunk with finish_reason
        done_chunk = {
            "id": request_id,
            "object": "chat.completion.chunk",
            "created": created,
            "model": model,
            "choices": [{
                "index": 0,
                "delta": {},
                "finish_reason": "stop",
            }],
        }
        if usage:
            done_chunk["usage"] = usage.model_dump()
        yield f"data: {_json.dumps(done_chunk)}\n\n"
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ─── Management Endpoints ────────────────────────────────────────────────────

@app.get("/gateway/stats")
async def get_stats(auth: bool = Depends(verify_auth)):
    """Get gateway statistics and KPIs."""
    summary = metrics.get_summary()
    budget = budget_guard.get_status()

    return {
        **summary,
        "budget": budget.model_dump(),
        "cache_stats": {
            "exact": exact_cache.get_stats(),
            "semantic": semantic_cache.get_stats(),
        },
        "free_api": {
            "success_today": int(metrics.get_daily_counter("free_api_success")),
            "failures_today": int(metrics.get_daily_counter("free_api_failure")),
            "escalations_avoided": int(metrics.get_daily_counter("free_api_success")),
        },
        "config": {
            "routing_strategy": config.routing_strategy,
            "mock_mode": config.mock_mode,
            "premium_model": config.providers.premium_model,
            "medium_model": config.providers.medium_model,
            "cheap_model": config.providers.cheap_model,
        },
    }


@app.post("/gateway/cache/invalidate")
async def invalidate_cache(event: CacheInvalidationEvent, auth: bool = Depends(verify_auth)):
    """Invalidate cache entries based on events."""
    if event.event == "git_commit":
        count = exact_cache.invalidate_by_type("code_suggestion")
        log.info(f"Cache invalidated for git_commit: {count} entries")
        return {"invalidated": True, "count": count, "reason": "git_commit"}

    elif event.event == "manual":
        exact_count = exact_cache.invalidate_all()
        semantic_count = semantic_cache.invalidate_all()
        return {"invalidated": True, "exact": exact_count, "semantic": semantic_count, "reason": "manual"}

    elif event.event == "file_change" and event.files:
        count = exact_cache.invalidate_by_type("code_suggestion")
        return {"invalidated": True, "count": count, "reason": "file_change"}

    return {"invalidated": False, "reason": "unknown_event"}


@app.post("/gateway/budget/kill")
async def activate_kill_switch(auth: bool = Depends(verify_auth)):
    """Manually activate the kill switch."""
    budget_guard.force_kill()
    return {"status": "kill_switch_activated", "premium_disabled": True}


@app.post("/gateway/budget/reset")
async def reset_kill_switch(auth: bool = Depends(verify_auth)):
    """Manually reset the kill switch."""
    budget_guard.force_reset()
    return {"status": "kill_switch_reset", "premium_disabled": False}


@app.get("/gateway/budget")
async def get_budget(auth: bool = Depends(verify_auth)):
    """Get current budget status."""
    return budget_guard.get_status().model_dump()


# ─── Dashboard ────────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def dashboard():
    """Serve the monitoring dashboard."""
    dashboard_path = Path(__file__).parent / "templates" / "dashboard.html"
    if dashboard_path.exists():
        return HTMLResponse(content=dashboard_path.read_text())
    # Fallback: redirect to chat console
    return HTMLResponse(content='<html><head><meta http-equiv="refresh" content="0;url=/chat"></head></html>')


@app.get("/chat", response_class=HTMLResponse)
async def chat_console():
    """Serve the interactive test console with image upload support."""
    chat_path = Path(__file__).parent / "templates" / "chat.html"
    if chat_path.exists():
        return HTMLResponse(content=chat_path.read_text())
    return HTMLResponse(content="""
    <html><body>
    <h1>Chat Console Not Found</h1>
    <p>Place chat.html in templates/ directory.</p>
    </body></html>
    """, status_code=404)


# ─── Debug View ──────────────────────────────────────────────────────────────

@app.get("/debug", response_class=HTMLResponse)
async def debug_console(request: Request):
    """Debug view — requires ?key= query param matching GATEWAY_SECRET."""
    key = request.query_params.get("key", "")
    if not key or not validate_api_key(key, gateway_secret):
        return HTMLResponse(content="<h1>401 — Add ?key=YOUR_GATEWAY_KEY</h1>", status_code=401)
    debug_path = Path(__file__).parent / "templates" / "debug.html"
    if debug_path.exists():
        return HTMLResponse(content=debug_path.read_text())
    return HTMLResponse(content="<h1>debug.html not found</h1>", status_code=404)


@app.get("/debug/data")
async def debug_data(auth: bool = Depends(verify_auth)):
    """Return debug buffer as JSON (newest first)."""
    return list(reversed(_debug_buffer))


# ─── File Extraction Endpoint ─────────────────────────────────────────────────

from pydantic import BaseModel as _PydBase

class FileExtractRequest(_PydBase):
    filename: str
    data_b64: str

@app.get("/v1/files/status")
async def file_extract_status(auth: bool = Depends(verify_auth)):
    """Check available file extraction libraries."""
    libs = {}
    for name, pkg in [("pymupdf", "fitz"), ("pdfplumber", "pdfplumber"),
                       ("pypdf", "pypdf"), ("python-docx", "docx"),
                       ("openpyxl", "openpyxl"), ("pytesseract", "pytesseract")]:
        try:
            mod = __import__(pkg)
            ver = getattr(mod, "__version__", getattr(mod, "VERSION", "installed"))
            libs[name] = str(ver)
        except ImportError:
            libs[name] = None
    return {"extractors": libs, "pdf_available": any(
        libs.get(l) for l in ["pymupdf", "pdfplumber", "pypdf"]
    )}


@app.post("/v1/files/extract")
async def extract_file(req: FileExtractRequest, auth: bool = Depends(verify_auth)):
    """
    Extract text from binary files (DOCX, XLSX, ZIP, etc.).
    Returns { text, type, pages?, files? }
    """
    import base64
    import tempfile

    ext = req.filename.rsplit(".", 1)[-1].lower() if "." in req.filename else ""
    raw = base64.b64decode(req.data_b64)

    try:
        # Audio files → redirect to transcription
        audio_exts = {"m4a", "mp3", "wav", "ogg", "webm", "flac", "aac", "wma", "opus"}
        if ext in audio_exts:
            mime_map = {
                "m4a": "audio/mp4", "mp3": "audio/mpeg", "wav": "audio/wav",
                "ogg": "audio/ogg", "webm": "audio/webm", "flac": "audio/flac",
                "aac": "audio/aac", "wma": "audio/x-ms-wma", "opus": "audio/opus",
            }
            result = await transcribe_audio(TranscribeRequest(
                audio_b64=req.data_b64,
                mime_type=mime_map.get(ext, "audio/mpeg"),
            ), auth=True)
            return {
                "text": result["text"],
                "type": "audio_transcription",
                "chars": len(result["text"]),
                "method": result.get("method", "unknown"),
            }

        if ext == "pdf":
            return await _extract_pdf(raw, req.filename)
        elif ext in ("docx", "doc"):
            return await _extract_docx(raw, req.filename)
        elif ext in ("xlsx", "xls", "csv"):
            return await _extract_xlsx(raw, req.filename, ext)
        elif ext in ("zip", "tar", "gz", "tgz"):
            return await _extract_archive(raw, req.filename, ext)
        else:
            # Try as plain text
            try:
                text = raw.decode("utf-8", errors="replace")
                return {"text": text, "type": "text", "chars": len(text)}
            except Exception:
                raise HTTPException(400, f"Unsupported file type: .{ext}")

    except HTTPException:
        raise
    except Exception as e:
        log.error(f"File extraction failed for {req.filename}: {e}")
        raise HTTPException(500, f"Failed to extract {req.filename}: {type(e).__name__}")


async def _extract_pdf(raw: bytes, filename: str) -> dict:
    """Extract text from PDF files. Tries multiple libraries with built-in fallback."""
    text = ""
    page_count = 0
    method = "none"

    # Try 1: pymupdf (fastest, best quality)
    try:
        import fitz
        doc = fitz.open(stream=raw, filetype="pdf")
        page_count = len(doc)
        pages = []
        for page in doc:
            page_text = page.get_text("text").strip()
            if page_text:
                pages.append(page_text)
        doc.close()
        text = "\n\n".join(pages)
        method = "pymupdf"
    except ImportError:
        log.warning("pymupdf not installed — pip install pymupdf --break-system-packages")
    except Exception as e:
        log.warning(f"pymupdf failed: {e}")

    # Try 2: pdfplumber
    if not text:
        try:
            import pdfplumber
            import io
            with pdfplumber.open(io.BytesIO(raw)) as pdf:
                page_count = len(pdf.pages)
                pages = []
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        pages.append(page_text.strip())
                    for table in page.extract_tables():
                        rows = [" | ".join(str(c or "") for c in row) for row in table]
                        pages.append("\n".join(rows))
                text = "\n\n".join(pages)
                method = "pdfplumber"
        except ImportError:
            log.warning("pdfplumber not installed")
        except Exception as e:
            log.warning(f"pdfplumber failed: {e}")

    # Try 3: pypdf
    if not text:
        try:
            from pypdf import PdfReader
            import io
            reader = PdfReader(io.BytesIO(raw))
            page_count = len(reader.pages)
            pages = [p.extract_text() or "" for p in reader.pages]
            text = "\n\n".join(p.strip() for p in pages if p.strip())
            method = "pypdf"
        except ImportError:
            log.warning("pypdf not installed")
        except Exception as e:
            log.warning(f"pypdf failed: {e}")

    # Try 4: Built-in raw text extraction (no dependencies!)
    if not text:
        try:
            import re
            # PDF text streams contain readable text between BT/ET markers
            # Also extract text from stream objects
            raw_text = raw.decode("latin-1", errors="replace")

            # Method A: Extract text between parentheses in text objects
            paren_texts = re.findall(r'\(([^)]{2,})\)', raw_text)
            # Filter out binary garbage
            readable = [t for t in paren_texts
                       if len(t) > 3 and sum(c.isalpha() or c.isspace() for c in t) > len(t) * 0.5]

            # Method B: Extract text from hex strings
            hex_texts = re.findall(r'<([0-9a-fA-F]{4,})>', raw_text)
            for h in hex_texts:
                try:
                    decoded = bytes.fromhex(h).decode("utf-16-be", errors="replace")
                    if len(decoded) > 2 and sum(c.isalpha() or c.isspace() for c in decoded) > len(decoded) * 0.3:
                        readable.append(decoded)
                except Exception:
                    pass

            if readable:
                # Deduplicate and join
                seen = set()
                unique = []
                for t in readable:
                    t = t.strip()
                    if t and t not in seen and len(t) > 2:
                        seen.add(t)
                        unique.append(t)
                text = " ".join(unique)
                method = "builtin-raw"
                page_count = raw_text.count("/Type /Page")
                log.info(f"PDF builtin extraction: {len(text)} chars from {len(unique)} fragments")
        except Exception as e:
            log.warning(f"Built-in PDF extraction failed: {e}")

    # Try 5: OCR as last resort (scanned PDFs)
    if not text or len(text.strip()) < 20:
        try:
            import fitz
            import io as _io
            import numpy as np
            doc = fitz.open(stream=raw, filetype="pdf")
            ocr_pages = []

            # Try PaddleOCR first
            paddle_model = None
            try:
                from vision_processor import _get_paddle_ocr
                paddle_model = _get_paddle_ocr()
            except Exception:
                pass

            for page in doc:
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                from PIL import Image
                img = Image.open(_io.BytesIO(img_bytes))

                if paddle_model:
                    # PaddleOCR path
                    img_array = np.array(img)
                    if hasattr(paddle_model, 'ocr'):
                        result = paddle_model.ocr(img_array, cls=True)
                    elif hasattr(paddle_model, 'predict'):
                        result = list(paddle_model.predict(img_array))
                    else:
                        result = None
                    page_lines = []
                    if result and result[0]:
                        for line in result[0]:
                            try:
                                if isinstance(line, (list, tuple)) and len(line) >= 2:
                                    text_info = line[1]
                                    if isinstance(text_info, (list, tuple)) and len(text_info) >= 2:
                                        t = str(text_info[0]).strip()
                                        if t:
                                            page_lines.append(t)
                                elif isinstance(line, dict):
                                    t = str(line.get('rec_text', line.get('text', ''))).strip()
                                    if t:
                                        page_lines.append(t)
                            except (ValueError, TypeError, IndexError):
                                continue
                    page_text = "\n".join(page_lines)
                else:
                    # Tesseract fallback
                    import pytesseract
                    page_text = pytesseract.image_to_string(img, lang="deu+eng").strip()

                if page_text:
                    ocr_pages.append(page_text)
            doc.close()
            if ocr_pages:
                text = "\n\n".join(ocr_pages)
                method = "paddleocr" if paddle_model else "tesseract_ocr"
        except Exception as ocr_err:
            log.warning(f"PDF OCR fallback failed: {ocr_err}")

    if not text.strip():
        # Give actionable error
        available = []
        for lib in ["fitz", "pdfplumber", "pypdf"]:
            try:
                __import__(lib)
                available.append(lib)
            except ImportError:
                pass
        if not available:
            raise HTTPException(
                501,
                "No PDF library installed! Run: pip install pymupdf --break-system-packages"
            )
        raise HTTPException(
            422,
            f"PDF appears empty or contains only images. "
            f"Available extractors: {', '.join(available)}"
        )

    # Truncate very long PDFs
    if len(text) > 80_000:
        text = text[:80_000] + f"\n\n[... truncated ...]"

    log.info(f"PDF extracted: {len(text)} chars, {page_count} pages, method={method}")
    return {
        "text": text,
        "type": "pdf",
        "pages": page_count,
        "chars": len(text),
        "method": method,
    }


async def _extract_docx(raw: bytes, filename: str) -> dict:
    """Extract text from DOCX files."""
    import tempfile
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            501, "python-docx not installed. Run: pip install python-docx --break-system-packages"
        )

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=True) as f:
        f.write(raw)
        f.flush()
        doc = Document(f.name)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract tables
    tables_text = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        tables_text.append("\n".join(rows))

    text = "\n".join(paragraphs)
    if tables_text:
        text += "\n\n" + "\n\n".join(tables_text)

    return {
        "text": text,
        "type": "docx",
        "paragraphs": len(paragraphs),
        "tables": len(doc.tables),
        "chars": len(text),
    }


async def _extract_xlsx(raw: bytes, filename: str, ext: str) -> dict:
    """Extract text from XLSX/CSV files."""
    import tempfile
    import csv
    import io

    if ext == "csv":
        text = raw.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = [" | ".join(row) for row in reader]
        return {"text": "\n".join(rows), "type": "csv", "rows": len(rows), "chars": len(text)}

    try:
        import openpyxl
    except ImportError:
        raise HTTPException(
            501, "openpyxl not installed. Run: pip install openpyxl --break-system-packages"
        )

    with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=True) as f:
        f.write(raw)
        f.flush()
        wb = openpyxl.load_workbook(f.name, read_only=True, data_only=True)

    all_text = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            all_text.append(f"[{sheet_name}]\n" + "\n".join(rows))

    wb.close()
    text = "\n\n".join(all_text)
    return {"text": text, "type": "xlsx", "sheets": len(wb.sheetnames), "chars": len(text)}


async def _extract_archive(raw: bytes, filename: str, ext: str) -> dict:
    """Extract and list contents of ZIP/TAR archives."""
    import tempfile
    import zipfile
    import tarfile
    import io

    extracted = []
    file_list = []

    if ext == "zip":
        with zipfile.ZipFile(io.BytesIO(raw)) as zf:
            for info in zf.infolist():
                if info.is_dir():
                    continue
                file_list.append(f"{info.filename} ({info.file_size:,} bytes)")
                # Extract text from small text files
                if info.file_size < 200_000:
                    fext = info.filename.rsplit(".", 1)[-1].lower() if "." in info.filename else ""
                    text_exts = {"txt","md","py","js","ts","json","xml","yaml","yml",
                                 "html","css","csv","sh","c","cpp","java","go","rs","rb","sql",
                                 "toml","ini","cfg","log","env"}
                    if fext in text_exts:
                        try:
                            content = zf.read(info.filename).decode("utf-8", errors="replace")
                            extracted.append(f"[{info.filename}]\n{content}")
                        except Exception:
                            pass

    elif ext in ("tar", "gz", "tgz"):
        mode = "r:gz" if ext in ("gz", "tgz") else "r:"
        try:
            with tarfile.open(fileobj=io.BytesIO(raw), mode=mode) as tf:
                for member in tf.getmembers():
                    if member.isdir():
                        continue
                    file_list.append(f"{member.name} ({member.size:,} bytes)")
                    if member.size < 200_000:
                        fext = member.name.rsplit(".", 1)[-1].lower() if "." in member.name else ""
                        text_exts = {"txt","md","py","js","ts","json","xml","yaml","yml",
                                     "html","css","csv","sh","toml","ini"}
                        if fext in text_exts:
                            try:
                                f = tf.extractfile(member)
                                if f:
                                    content = f.read().decode("utf-8", errors="replace")
                                    extracted.append(f"[{member.name}]\n{content}")
                            except Exception:
                                pass
        except Exception:
            raise HTTPException(400, "Could not read archive. Unsupported format.")

    text = f"Archive: {filename}\nFiles: {len(file_list)}\n\n"
    text += "File listing:\n" + "\n".join(file_list)
    if extracted:
        text += "\n\n" + "\n\n".join(extracted)

    # Truncate if too large (keep first 50k chars)
    if len(text) > 50_000:
        text = text[:50_000] + f"\n\n[... truncated, {len(text)-50_000} chars remaining ...]"

    return {"text": text, "type": "archive", "files": len(file_list), "chars": len(text)}


# ─── Audio Transcription Endpoint (2-tier: Groq API → local Whisper) ──────────

class TranscribeRequest(_PydBase):
    audio_b64: str  # Base64-encoded audio data
    mime_type: str = "audio/webm"  # Browser MediaRecorder default
    language: str = ""  # Optional: "de", "en", etc. Empty = auto-detect

# Lazy-loaded local whisper model (loads on first use)
_local_whisper_model = None
_local_whisper_lock = None


def _get_local_whisper():
    """Lazy-load faster-whisper model. Downloads ~400MB on first use."""
    global _local_whisper_model, _local_whisper_lock
    import threading
    if _local_whisper_lock is None:
        _local_whisper_lock = threading.Lock()

    if _local_whisper_model is not None:
        return _local_whisper_model

    with _local_whisper_lock:
        if _local_whisper_model is not None:
            return _local_whisper_model
        try:
            from faster_whisper import WhisperModel
            # "base" = 140MB, good quality/speed tradeoff on CPU
            # Options: tiny (75MB, fast), base (140MB), small (460MB, best quality)
            model_size = os.environ.get("WHISPER_MODEL", "base")
            log.info(f"Loading local Whisper model: {model_size} (first load downloads the model)...")
            _local_whisper_model = WhisperModel(
                model_size,
                device="cpu",
                compute_type="int8",  # Fastest on CPU
                cpu_threads=min(4, os.cpu_count() or 2),
            )
            log.info(f"Local Whisper model loaded: {model_size}")
            return _local_whisper_model
        except ImportError:
            log.warning("faster-whisper not installed: pip install faster-whisper --break-system-packages")
            return None
        except Exception as e:
            log.error(f"Failed to load local Whisper: {e}")
            return None


@app.get("/v1/audio/status")
async def audio_status(auth: bool = Depends(verify_auth)):
    """Check available transcription backends."""
    groq_ok = bool(os.environ.get("GROQ_API_KEY", ""))
    local_ok = False
    local_model = None
    try:
        import faster_whisper
        local_ok = True
        local_model = os.environ.get("WHISPER_MODEL", "base")
    except ImportError:
        pass
    browser_stt = True  # Always available (client-side)
    return {
        "browser_speech_api": browser_stt,
        "groq_whisper": groq_ok,
        "local_whisper": local_ok,
        "local_model": local_model,
        "available": browser_stt or groq_ok or local_ok,
    }


async def _transcribe_internal(audio_bytes: bytes, mime_type: str = "audio/ogg",
                                language: str = "de") -> dict:
    """
    Internal transcription helper (no auth). Used by:
    - Chat pipeline auto-transcription (audio in messages)
    - /v1/audio/transcribe endpoint
    Returns {"text": "...", "method": "groq"|"local", "duration_ms": N}
    """
    import tempfile

    if len(audio_bytes) < 100:
        return {"text": "", "method": "skipped", "error": "audio too short"}
    if len(audio_bytes) > 25 * 1024 * 1024:
        return {"text": "", "method": "skipped", "error": "audio too large"}

    ext_map = {
        "audio/webm": ".webm", "audio/mp4": ".m4a", "audio/mpeg": ".mp3",
        "audio/wav": ".wav", "audio/ogg": ".ogg", "audio/flac": ".flac",
        "audio/opus": ".opus",
    }
    ext = ext_map.get(mime_type, ".ogg")

    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as f:
        f.write(audio_bytes)
        tmp_path = f.name

    try:
        # Tier 1: Groq
        groq_key = os.environ.get("GROQ_API_KEY", "")
        if groq_key:
            try:
                text, latency = await _transcribe_groq(tmp_path, ext, mime_type, language, groq_key)
                log.info(f"Auto-transcription [groq]: {len(audio_bytes)}B → {len(text)} chars, "
                         f"{latency:.0f}ms")
                return {"text": text, "method": "groq", "duration_ms": round(latency)}
            except Exception as e:
                log.warning(f"Groq auto-transcription failed: {e}, trying local")

        # Tier 2: Local whisper
        model = _get_local_whisper()
        if model:
            try:
                text, latency = await _transcribe_local(model, tmp_path, language)
                log.info(f"Auto-transcription [local]: {len(audio_bytes)}B → {len(text)} chars, "
                         f"{latency:.0f}ms")
                return {"text": text, "method": "local", "duration_ms": round(latency)}
            except Exception as e:
                log.error(f"Local auto-transcription failed: {e}")
                return {"text": "", "method": "failed", "error": str(e)}

        return {"text": "", "method": "unavailable", "error": "no transcription backend"}
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


@app.post("/v1/audio/transcribe")
async def transcribe_audio(req: TranscribeRequest, auth: bool = Depends(verify_auth)):
    """
    Transcribe audio to text.
    Tier 1: Groq Whisper API (fast, free tier)
    Tier 2: Local faster-whisper (offline, no API key needed)
    """
    import tempfile

    # Decode audio
    try:
        audio_bytes = base64.b64decode(req.audio_b64)
    except Exception:
        raise HTTPException(400, "Invalid base64 audio data")

    if len(audio_bytes) < 100:
        raise HTTPException(400, "Audio too short")
    if len(audio_bytes) > 25 * 1024 * 1024:  # 25MB limit
        raise HTTPException(413, "Audio too large (max 25MB)")

    # Determine file extension from mime type
    ext_map = {
        "audio/webm": ".webm",
        "audio/mp4": ".m4a",
        "audio/mpeg": ".mp3",
        "audio/wav": ".wav",
        "audio/ogg": ".ogg",
        "audio/flac": ".flac",
    }
    ext = ext_map.get(req.mime_type, ".webm")

    # Write to temp file
    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as f:
        f.write(audio_bytes)
        tmp_path = f.name

    try:
        # ── Tier 1: Groq Whisper API (fast, free) ──
        groq_key = os.environ.get("GROQ_API_KEY", "")
        if groq_key:
            try:
                text, latency = await _transcribe_groq(tmp_path, ext, req.mime_type,
                                                        req.language, groq_key)
                log.info(f"Transcription [groq]: {len(audio_bytes)}B → {len(text)} chars, "
                         f"{latency:.0f}ms, lang={req.language or 'auto'}")
                return {
                    "text": text,
                    "method": "groq",
                    "duration_ms": round(latency),
                    "audio_size": len(audio_bytes),
                }
            except Exception as e:
                log.warning(f"Groq Whisper failed: {e}, trying local fallback")

        # ── Tier 2: Local faster-whisper (offline) ──
        model = _get_local_whisper()
        if model:
            try:
                text, latency = await _transcribe_local(model, tmp_path, req.language)
                log.info(f"Transcription [local]: {len(audio_bytes)}B → {len(text)} chars, "
                         f"{latency:.0f}ms, lang={req.language or 'auto'}")
                return {
                    "text": text,
                    "method": "local",
                    "duration_ms": round(latency),
                    "audio_size": len(audio_bytes),
                }
            except Exception as e:
                log.error(f"Local Whisper failed: {e}")
                raise HTTPException(500, f"Local transcription failed: {type(e).__name__}")

        # No backend available
        raise HTTPException(
            501,
            "Audio transcription not available. "
            "Install on server: pip install faster-whisper --break-system-packages"
        )

    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


async def _transcribe_groq(tmp_path: str, ext: str, mime_type: str,
                            language: str, api_key: str) -> tuple[str, float]:
    """Transcribe via Groq Whisper API."""
    start = time.time()
    async with httpx.AsyncClient() as client:
        with open(tmp_path, "rb") as audio_file:
            files = {"file": (f"audio{ext}", audio_file, mime_type)}
            data = {
                "model": "whisper-large-v3-turbo",
                "response_format": "json",
            }
            if language:
                data["language"] = language

            response = await client.post(
                "https://api.groq.com/openai/v1/audio/transcriptions",
                headers={"Authorization": f"Bearer {api_key}"},
                files=files,
                data=data,
                timeout=30.0,
            )

    if response.status_code != 200:
        raise Exception(f"Groq API {response.status_code}: {response.text[:200]}")

    result = response.json()
    text = result.get("text", "").strip()
    latency = (time.time() - start) * 1000
    return text, latency


async def _transcribe_local(model, tmp_path: str, language: str) -> tuple[str, float]:
    """Transcribe via local faster-whisper model."""
    import asyncio
    start = time.time()

    def _run():
        kwargs = {"beam_size": 1, "vad_filter": True}
        if language:
            kwargs["language"] = language
        segments, info = model.transcribe(tmp_path, **kwargs)
        text = " ".join(seg.text.strip() for seg in segments)
        return text, info

    # Run in thread pool to not block the event loop
    loop = asyncio.get_event_loop()
    text, info = await loop.run_in_executor(None, _run)
    latency = (time.time() - start) * 1000
    detected_lang = getattr(info, 'language', '?')
    log.debug(f"Local Whisper: detected_lang={detected_lang}, "
              f"lang_prob={getattr(info, 'language_probability', 0):.2f}")
    return text, latency


# ─── Run ──────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
